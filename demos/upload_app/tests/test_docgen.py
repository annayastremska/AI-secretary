# -*- coding: utf-8 -*-
"""Режим створення документа: чат питає поле за полем.

Критерії приймання — `docs/tasks/2026-08-27_acceptance-criteria.md`, розділ 12.
Пропозиція й розбір ризиків — `docs/tasks/2026-08-28_dovidka-mode-proposal.md`.
Тести написані ДО коду.

Головне, що вони охороняють, — те, на чому режим ламався б молча: **жодне
значення бланка не лишається порожнім тихо**. Схема, якою система документ
ЧИТАЄ, і заповнювач, який його ПИШЕ, говорять різними словниками (заміряно:
спільних імен нуль), і між ними переклад із похідними — ініціали, число
словами, дата частинами, формулювання, залежні від статі. Тому перелік питань
виводиться з того, чого чекає ЗАПОВНЮВАЧ, а схема лишається джерелом перевірок.
"""
import os
import re
import sys

import pytest

HERE = os.path.dirname(os.path.abspath(__file__))
APP_DIR = os.path.dirname(HERE)
for p in (APP_DIR, os.path.join(APP_DIR, "chat_gradio")):
    if p not in sys.path:
        sys.path.insert(0, p)

from chat_gradio import docgen                 # noqa: E402
from chat_gradio import tiers as tier_chat     # noqa: E402


# ── Види документів і план питань ────────────────────────────────────────────


def test_two_kinds_and_both_blanks_exist():
    ks = docgen.kinds()
    assert {k["key"] for k in ks} == {"leave", "deployment"}
    for k in ks:
        assert k["title"].strip()
        assert os.path.exists(k["blank"]), k["blank"]


@pytest.mark.parametrize("kind", ["leave", "deployment"])
def test_plan_has_asks_and_kinds(kind):
    fields = docgen.plan(kind)
    assert fields, kind
    for f in fields:
        assert f["ask"].strip().endswith((".", "?", ":")), f
        assert f["type"] in ("person", "date", "text", "number"), f
        assert isinstance(f["required"], bool)


@pytest.mark.parametrize("kind,need", [
    ("leave", {"person", "leave_type", "place", "start", "end", "number"}),
    ("deployment", {"person", "dest", "dest_org", "purpose", "start", "end",
                    "number", "order_number"}),
])
def test_required_fields_are_asked(kind, need):
    names = {f["name"] for f in docgen.plan(kind) if f["required"]}
    assert need <= names, need - names


# ── К2: НІЧОГО не лишається порожнім тихо ───────────────────────────────────


@pytest.mark.parametrize("kind", ["leave", "deployment"])
def test_every_blank_value_is_covered(kind):
    """Головний тест мапінгу.

    Перелік ключів, яких чекає заповнювач, беремо з НЬОГО САМОГО (розбором
    його коду), а не з нашого уявлення про нього. Кожен ключ мусить бути або
    питаним, або похідним, або константою — інакше у файлі лишиться порожнє
    місце, а людина відповіла на всі питання й вважає, що документ готовий.
    """
    needed = docgen.blank_keys(kind)
    assert needed, kind
    covered = docgen.covered_keys(kind)
    missing = sorted(needed - covered)
    assert not missing, (
        f"{kind}: ці значення бланка нічим не заповнюються: {missing}")


# ── Особа: лише зі штатки ───────────────────────────────────────────────────


def test_person_found_by_full_name():
    """Прізвище «Гавриш» у штатці НЕ унікальне (їх двоє) -- і це не дефект, а
    життя. Унікальність перевіряємо повним ПІБ."""
    rows, err = docgen.find_person("Гавриш Адам Станіславович")
    assert err is None, err
    assert len(rows) == 1
    assert rows[0]["last_name"] == "Гавриш"
    assert rows[0]["gender"] in ("чоловіча", "жіноча")


def test_ambiguous_surname_is_asked_again():
    rows, err = docgen.find_person("Гавриш")
    assert len(rows) >= 2, rows
    assert err and "кого саме" in err.lower(), err


def test_person_found_by_full_name_in_any_case():
    rows, err = docgen.find_person("гавриш адам станіславович")
    assert err is None and len(rows) == 1


def test_unknown_person_is_refused():
    rows, err = docgen.find_person("Кривопишний Аристарх")
    assert not rows
    assert "штатці" in (err or ""), err


def test_several_matches_are_listed_not_guessed():
    """К4: два Приймаки — це питання до людини, а не вибір за неї."""
    rows, err = docgen.find_person("Приймак")
    if len(rows) < 2:
        pytest.skip("у штатці немає однофамільців для цієї перевірки")
    assert err and "кого саме" in err.lower(), err
    for r in rows[:3]:
        assert r["full_name"] in err


# ── Дати, тривалість, числа ─────────────────────────────────────────────────


def test_date_parsed_and_normalised():
    val, err, note = docgen.validate_answer(
        {"name": "start", "type": "date", "required": True, "ask": "Дата."},
        "21 вересня 2026", {})
    assert err is None, err
    assert val == "2026-09-21", val


def test_impossible_date_refused():
    val, err, _ = docgen.validate_answer(
        {"name": "start", "type": "date", "required": True, "ask": "Дата."},
        "31 лютого 2026", {})
    assert val is None and "не існує" in (err or "").lower(), err


def test_end_before_start_refused():
    val, err, _ = docgen.validate_answer(
        {"name": "end", "type": "date", "required": True, "ask": "Дата."},
        "2026-09-01", {"start": "2026-09-10"})
    assert val is None and err, err
    assert "рані" in err.lower()


def test_duration_is_computed_and_named():
    val, err, note = docgen.validate_answer(
        {"name": "end", "type": "date", "required": True, "ask": "Дата."},
        "2026-10-10", {"start": "2026-09-21"})
    assert err is None
    assert "20" in (note or ""), note


def test_number_must_be_digits():
    val, err, _ = docgen.validate_answer(
        {"name": "number", "type": "number", "required": True, "ask": "Номер."},
        "№3О4", {})
    assert val is None and err


# ── Збирання: скасування, пропуск, три спроби ───────────────────────────────


def test_cancel_exits_at_any_step():
    st = docgen.start()
    st, reply, path = docgen.step(st, "відпускний квиток")
    st, reply, path = docgen.step(st, "скасувати")
    assert path is None
    assert docgen.is_done(st)
    assert "скасовано" in reply.lower()


def test_skip_only_for_optional_fields():
    st = docgen.start()
    st, _, _ = docgen.step(st, "відпускний квиток")
    # Перше поле обовʼязкове -- пропустити не можна.
    st, reply, _ = docgen.step(st, "пропустити")
    assert "обов" in reply.lower(), reply


def test_three_failed_tries_offer_a_way_out():
    """К14: не крутимо те саме питання далі -- інакше людина в пастці."""
    st = docgen.start()
    st, _, _ = docgen.step(st, "відпускний квиток")
    for _ in range(3):
        st, reply, _ = docgen.step(st, "щось незрозуміле")
    assert "скасувати" in reply.lower() or "пропустити" in reply.lower(), reply


def test_ordinary_question_is_not_stored_as_a_value():
    """К15: у режимі не пишемо в поле те, що явно є питанням до бази."""
    st = docgen.start()
    st, _, _ = docgen.step(st, "відпускний квиток")
    st, reply, _ = docgen.step(st, "а скільки зараз у відпустці?")
    assert st["answers"].get("person") is None
    assert "зараз збирається документ" in reply.lower(), reply


def test_unknown_kind_is_asked_again():
    st = docgen.start()
    st, reply, _ = docgen.step(st, "довідка про доходи")
    assert st.get("kind") is None
    assert "відпускний" in reply.lower()


# ── Наскрізь по живій базі ──────────────────────────────────────────────────


def _db_reachable():
    try:
        with tier_chat._connect() as conn, conn.cursor() as cur:
            cur.execute("SELECT 1")
            return True
    except Exception:
        return False


db_only = pytest.mark.skipif(not _db_reachable(), reason="потрібна жива база")


@db_only
def test_taken_numbers_and_next_free():
    taken = docgen.taken_numbers()
    assert taken, "у базі мусять бути номери документів"
    free = docgen.next_free_number(taken)
    assert free not in taken


@db_only
def test_conflict_is_found_for_the_same_person():
    """К8: перетин у ТІЄЇ САМОЇ особи. Різні люди одночасно — норма."""
    # Ґоляш: відпустка 2026-09-21 — 2026-10-10 (документ №1077).
    rows, err = docgen.find_person("Ґоляш")
    assert not err and rows, err
    sid = rows[0]["service_id"]
    clash = docgen.find_conflicts(sid, "leave", "2026-10-01", "2026-10-05")
    assert clash, "перетин мусив знайтись"
    assert any("1077" in str(c.get("number") or "") for c in clash), clash


@db_only
def test_touching_days_are_not_a_conflict():
    rows, _ = docgen.find_person("Ґоляш")
    sid = rows[0]["service_id"]
    # Наступний день після завершення -- не перетин.
    assert docgen.find_conflicts(sid, "leave", "2026-10-11", "2026-10-20") == []


@db_only
def test_leave_and_deployment_overlap_is_also_a_conflict():
    """ОБИДВА види відсутності разом (Аня 28.08).

    Цей тест раніше стверджував протилежне -- «різні виміри, перетин
    шукається в межах одного» -- і тим закріплював дірку: фізично людина не
    може бути одночасно у відпустці й у відрядженні. Мій власний тест
    описував не правило продукту, а те, як я випадково написала запит.

    Ґоляш: відпустка 2026-09-21 — 2026-10-10. Відрядження всередині цього
    періоду -- суперечність.
    """
    rows, _ = docgen.find_person("Ґоляш")
    sid = rows[0]["service_id"]
    clash = docgen.find_conflicts(sid, "deployment", "2026-10-01", "2026-10-05")
    assert clash, "перетин відрядження з відпусткою мусить ловитись"
    assert any("1077" in str(c.get("number") or "") for c in clash), clash


@db_only
def test_conflict_names_the_kind_of_the_existing_document():
    """Людина мусить бачити, ЩО саме перетинається: інакше «вже є документ
    №1077» на питання про відрядження читається як помилка системи."""
    rows, _ = docgen.find_person("Ґоляш")
    sid = rows[0]["service_id"]
    clash = docgen.find_conflicts(sid, "deployment", "2026-10-01", "2026-10-05")
    assert clash[0].get("kind_label") == "відпустка", clash[0]


@db_only
def test_still_no_conflict_outside_the_period():
    rows, _ = docgen.find_person("Ґоляш")
    sid = rows[0]["service_id"]
    assert docgen.find_conflicts(sid, "deployment", "2026-11-01",
                                 "2026-11-05") == []


# ── Побудова файла ──────────────────────────────────────────────────────────


ANSWERS_LEAVE = {
    "person": "UNIT-0001",
    "leave_type": "щорічна основна відпустка за 2026 рік",
    "place": "м. Кривоярськ",
    "start": "2026-11-03",
    "end": "2026-11-12",
    "number": "9001",
    "issue": "2026-11-01",
}


def test_build_makes_a_file_with_the_mark(tmp_path):
    """К12: позначку «сформовано автоматично» прибрати неможливо."""
    path = docgen.build("leave", ANSWERS_LEAVE, out_dir=str(tmp_path))
    assert os.path.exists(path) and path.endswith(".docx")
    from docx import Document
    text = "\n".join(p.text for p in Document(path).paragraphs)
    for t in Document(path).tables:
        for row in t.rows:
            text += "\n" + " ".join(c.text for c in row.cells)
    assert "сформовано автоматично" in text.lower(), text[:400]


def test_build_fills_what_was_answered(tmp_path):
    path = docgen.build("leave", ANSWERS_LEAVE, out_dir=str(tmp_path))
    from docx import Document
    text = "\n".join(p.text for p in Document(path).paragraphs)
    for t in Document(path).tables:
        for row in t.rows:
            text += "\n" + " ".join(c.text for c in row.cells)
    assert "9001" in text
    assert "Кривоярськ" in text
    assert "ГАВРИШ" in text.upper()


@db_only
def test_build_writes_nothing_to_the_database(tmp_path):
    """К13: файл — не запис. Кількість документів і фактів не змінюється."""
    def counts():
        with tier_chat._connect() as conn, conn.cursor() as cur:
            cur.execute("SELECT count(*) AS n FROM documents")
            docs = cur.fetchone()["n"]
            cur.execute("SELECT count(*) AS n FROM facts")
            return docs, cur.fetchone()["n"]

    before = counts()
    docgen.build("leave", ANSWERS_LEAVE, out_dir=str(tmp_path))
    assert counts() == before


def test_temp_files_are_cleaned(tmp_path):
    """Ш6: тимчасові файли не накопичуються."""
    old = tmp_path / "старий.docx"
    old.write_bytes(b"x")
    os.utime(old, (0, 0))
    docgen.build("leave", ANSWERS_LEAVE, out_dir=str(tmp_path))
    assert not old.exists(), "старий файл мусив прибратись"


# ── Ш3: важкі бібліотеки не тягнуться при імпорті ───────────────────────────


def test_heavy_imports_are_lazy():
    """Ш3: PIL і numpy тягне генератор демо-набору. У веб-процесі вони не
    мусять з'являтись від самого імпорту режиму."""
    src = open(os.path.join(APP_DIR, "chat_gradio", "docgen.py"),
               encoding="utf-8").read()
    head = src.split("def ", 1)[0]
    for heavy in ("import numpy", "from PIL", "import PIL"):
        assert heavy not in head, heavy


# ── К1, К16: підключення до чата ─────────────────────────────────────────────


def _app_src():
    return open(os.path.join(APP_DIR, "chat_gradio", "app.py"),
                encoding="utf-8").read()


def test_button_is_under_clear_chat():
    """К1: кнопка стоїть саме під «Очистити чат», як просила Аня."""
    src = _app_src()
    at_clear = src.index('gr.Button("Очистити чат"')
    at_docgen = src.index('gr.Button("Створити документ"')
    assert at_clear < at_docgen
    # І між ними нічого іншого не встряло.
    between = src[at_clear:at_docgen]
    assert "gr.Button(" not in between.split('elem_id="new-chat")', 1)[1]


def test_mode_state_lives_in_the_session():
    """К16: `gr.State`, не модульна змінна.

    Модульна дала б «один увімкнув -- у всіх увімкнулось»: найтиповіша помилка
    такого перемикача, і на демо вона виглядала б як зламаний чат.
    """
    src = _app_src()
    assert "docgen_state = gr.State(None)" in src
    # У самому модулі режиму глобального стану немає.
    dg = open(os.path.join(APP_DIR, "chat_gradio", "docgen.py"),
              encoding="utf-8").read()
    head = dg.split("def ", 1)[0]
    for bad in ("STATE =", "MODE =", "_STATE =", "SESSION ="):
        assert bad not in head, bad


def test_respond_takes_and_returns_the_mode_state():
    src = _app_src()
    assert "def respond(message, history, dg=None):" in src
    assert "outs = [box, chat, hero, retry, last_q, docgen_state, docfile]" in src


def test_reset_clears_the_mode():
    """Ш7: «Очистити чат» мусить вийти з режиму, інакше людина лишиться в
    ньому після очищення й не зрозуміє, чому чат не відповідає на питання."""
    src = _app_src()
    reset = src.split("def reset():", 1)[1].split("\n    #", 1)[0]
    assert "None, gr.update(value=None, visible=False)" in reset


def test_mode_errors_do_not_break_the_chat():
    """Режим -- додаткова можливість. Виняток у ньому не має валити хід."""
    src = _app_src()
    branch = src.split("if dg and not docgen.is_done(dg):", 1)[1][:900]
    assert "except Exception" in branch
    assert "Створення документа перервано" in branch


# ── Перетин: підстава, наполягання, вихід ────────────────────────────────────


def test_override_outside_the_conflict_is_not_data():
    """Керівна фраза, що стала значенням поля, -- найгірший вид тихої помилки:
    людина думає, що керує, а вона диктує.

    У живому прогоні «формуй усе одно» записалось у номер перевізного
    документа, бо перетину в той момент не було.
    """
    st = docgen.start()
    st, _, _ = docgen.step(st, "відпускний квиток")
    st, reply, _ = docgen.step(st, "формуй усе одно")
    assert st["answers"].get("person") is None
    assert "перетин" in reply.lower(), reply


def test_conflict_check_runs_as_soon_as_dates_are_known():
    """Умова мусить залежати від ДАНИХ, а не від позиції в переліку.

    Спершу перевірка стояла «перед питанням про номер» -- а в квитку між
    датами й номером є ще ВПД і супутники, тобто людина відповідала на два
    зайві питання, перш ніж дізнатись, що документ не потрібен.
    """
    src = open(os.path.join(APP_DIR, "chat_gradio", "docgen.py"),
               encoding="utf-8").read()
    block = src.split("def _advance", 1)[1][:900]
    assert 'a.get("person") and a.get("start") and a.get("end")' in block
    assert 'nxt["name"] == "number"' not in block.split("conflict_checked")[0]


@db_only
def test_conflict_stops_the_document_and_offers_the_legal_route():
    st = docgen.start()
    for a in ("відпускний квиток", "Ґоляш Богодар Святославович",
              "щорічна основна відпустка за 2026 рік", "м. Сухобрід",
              "1 жовтня 2026", "5 жовтня 2026"):
        st, reply, path = docgen.step(st, a)
    assert path is None
    assert st.get("conflict"), reply
    assert "перетина" in reply.lower(), reply
    assert "замість" in reply.lower(), reply
    # Підстава -- і збирання продовжується.
    st, reply, path = docgen.step(st, "замість №1077")
    assert not st.get("conflict")
    assert "1077" in reply


# ── Номер звернення й слід режиму ────────────────────────────────────────────


def test_session_gets_a_request_id():
    st = docgen.start()
    assert re.fullmatch(r"[0-9a-f]{6}", st["id"]), st.get("id")


def test_id_is_shown_at_the_start_and_with_the_file(tmp_path):
    st = docgen.start()
    st, reply, _ = docgen.step(st, "")
    assert st["id"] in reply, reply
    # І в кінці, разом із файлом.
    line = docgen.done_line(st)
    assert st["id"] in line


def test_trace_file_is_separate_from_the_chat_trace():
    """Змішувати не можна, і це не смак.

    `trace_lookup --check` судить кожен запис за правилом «у відповіді є
    джерело». У режимі створення відповідь -- це ПИТАННЯ до людини, джерела в
    неї немає за побудовою, тому в спільному файлі прилад почав би рахувати
    порушення там, де їх нема. Ми вже наступали на це, коли тести писали в
    бойовий слід.
    """
    from chat_gradio import trace as chat_trace
    assert docgen.TRACE_PATH != chat_trace.TRACE_PATH
    assert "docgen" in os.path.basename(docgen.TRACE_PATH)


def test_trace_records_no_personal_data(tmp_path, monkeypatch):
    """Правило проєкту: у логи не їде персональна інформація.

    Тут це важливіше, ніж у чаті: у режимі людина ДИКТУЄ ПІБ, дати й місце.
    Тому в слід ідуть лише назва поля, чи воно прийнялось, і чи був перетин --
    без жодного значення.
    """
    path = tmp_path / "docgen-trace.jsonl"
    monkeypatch.setattr(docgen, "TRACE_PATH", str(path))
    st = docgen.start()
    st, _, _ = docgen.step(st, "відпускний квиток")
    st, _, _ = docgen.step(st, "Гавриш Адам Станіславович")
    body = path.read_text(encoding="utf-8")
    assert "Гавриш" not in body, body
    assert "UNIT-" not in body, body
    assert st["id"] in body
    assert "person" in body


# ── Формат відповіді: очікуваний вигляд і ОДИН приклад ──────────────────────


@pytest.mark.parametrize("kind", ["leave", "deployment"])
def test_text_fields_state_the_expected_format_with_one_example(kind):
    """Рішення Ані: за людину не нормалізуємо, але мусимо сказати, ЯК саме
    писати, і дати один приклад."""
    for f in docgen.plan(kind):
        if f["type"] != "text":
            continue
        ask = f["ask"]
        assert "наприклад" in ask.lower(), (kind, f["name"], ask)
        # Один приклад, не три: перелік прикладів читається як перелік
        # варіантів на вибір.
        assert ask.lower().count("наприклад") == 1, (kind, f["name"], ask)


def test_summary_shows_the_person_by_name_not_by_code():
    """У зведенні стояло «Прізвище, ім'я, по батькові: UNIT-0026».

    Службовий код у полі, яке людина щойно назвала словами, -- це та сама
    внутрішня кухня на екрані, на яку скаржився Денис (п. 25).
    """
    st = docgen.start()
    st["kind"] = "leave"
    st["answers"] = {"person": "UNIT-0001", "leave_type": "щорічна",
                     "place": "м. Рівне", "start": "2026-11-03",
                     "end": "2026-11-12", "number": "9001",
                     "issue": "2026-11-01"}
    text = docgen._summary(st)
    assert "UNIT-0001" not in text, text
    assert "Гавриш" in text, text


def test_request_id_is_not_duplicated_on_screen():
    """Аня побачила «звернення bcf37d» двічі в одному повідомленні.

    Причина: `render_reply` виносить номер у дрібний підпис знизу, а з тіла
    вирізає його регуляркою, де ДВОКРАПКА стоїть у шаблоні. Мій рядок був без
    двокрапки, тому не вирізався -- і номер лишався в тексті плюс додавався
    підписом. Формат рядка тут -- частина стику з рендером, а не оформлення.
    """
    from chat_gradio import app as chat_app
    st = docgen.start()
    st, reply, _ = docgen.step(st, "")
    out = chat_app.render_reply(reply)
    assert out.count(st["id"]) == 1, out
    assert 'class="req-id"' in out, out
