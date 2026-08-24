"""Точка входу пайплайна: перетворює файл на (ocr_text, ocr_blocks) --
той самий формат, що очікує extract_document(), незалежно від джерела.

docx має вбудований текстовий шар -> пряма екстракція (без OCR).
Зображення (скан/фото) -> OCR (Surya, підключається ззовні через ocr_fn,
щоб цей модуль не тягнув важку ML-залежність, коли вона не потрібна).
"""
import hashlib
import os

DOCX_EXTS = (".docx",)
PDF_EXTS = (".pdf",)
IMAGE_EXTS = (".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tif", ".tiff")

# Нижче цього середнього числа символів на сторінку вважаємо, що текстового
# шару фактично немає (скан, загорнутий у PDF) -- тоді сторінки рендеряться
# в зображення й ідуть в OCR. Поріг свідомо низький: краще спробувати
# текстовий шар, ніж марно ганяти OCR по нормальному PDF.
# Нижні межі ПРАВДОПОДІБНОГО виходу OCR для аркуша бланка. Не критерій
# правильності й не поріг статусу -- лише привід сказати вголос "схоже, OCR
# деградував". Числа взяті з заміру на реальних фото набору: успішне
# розпізнавання дає 25+ блоків і 900+ символів, тому 5 і 200 -- це з великим
# запасом нижче норми, а не підігнана межа.
MIN_PLAUSIBLE_OCR_BLOCKS = 5
MIN_PLAUSIBLE_OCR_CHARS = 200

MIN_TEXT_CHARS_PER_PAGE = 40
PDF_RENDER_DPI = 200

#: Позначка ПОХОДЖЕННЯ блоку: текстовий шар PDF (`page.get_text("blocks")`).
#: Потрібна тому, що "блок" у трьох джерел означає РІЗНЕ, і надійність межі
#: блоку теж різна:
#:   docx  -- абзац або клітинка, тобто справжня межа поля;
#:   Surya -- область зображення (межі відновлює blank_form.resegment_by_blank);
#:   PDF   -- ГРУПА послідовних рядків, яку PyMuPDF склеїла за близькістю.
#: Останнє -- не абзац у жодному сенсі: один блок може містити кілька полів
#: бланка підряд (звідси `value_starts_after` і `_value_lines_after_label_note`
#: в extract.py), а ОДНЕ значення, розбите переносом рядка, може бути
#: розкладене на два блоки (розд. 5.7, `position_and_workplace` на pdf).
#: Тому саме на цьому шляху межа блоку НЕ є межею поля, і extract.py має право
#: добирати хвіст значення через неї -- але тільки тут, щоб docx і фото
#: лишались недоторканими за побудовою.
PDF_TEXT_SOURCE = "pdf_text"


def extract_docx_blocks(path: str):
    """Абзаци + клітинки таблиць як блоки, у порядку появи в документі.
    Розбиття багаторядкових блоків на окремі рядки -- відповідальність
    extract_document()/flatten_blocks() (той самий дефект гранулярності
    трапляється і в блоках Surya, тому виправлено централізовано там, а не
    тут по джерелу).

    Імпорт python-docx -- лінивий, усередині функції: цей модуль
    використовується і для суто-OCR прогонів (зображення), де docx не
    встановлено й не потрібен -- те саме міркування, що для Surya (ocr_fn
    підключається ззовні саме щоб не тягнути важку залежність завжди)."""
    from docx import Document

    doc = Document(path)
    blocks = []

    # Об'єднана клітинка (merged) повертається python-docx ОКРЕМО для кожної
    # колонки сітки, під якою вона лежить, тому наївний обхід row.cells
    # дублював той самий текст 4-6 разів (підтверджено на реальному
    # "Посвідченні про відрядження": "рядовий БЕВЗЕНКО ..." з'являвся
    # чотири рази). Це не косметика: роздутий текст -- це довший промпт для
    # LLM, тобто прямі витрати часу на CPU. Дедуплікуємо за самим XML-
    # елементом клітинки (cell._tc), а не за текстом: дві РІЗНІ клітинки з
    # однаковим текстом (напр. дві порожні дати) лишаються обидві.
    # lxml перевикористовує id() звільнених проксі-об'єктів, тому набору
    # самих id НЕ достатньо: перший варіант цього коду через id() випадково
    # відкинув рядок з датами відрядження (елемент вважався "вже баченим").
    # Тримаємо сильні посилання на елементи -- поки посилання живе, lxml
    # гарантує той самий проксі, отже id стабільний і не перевикористаний.
    # Визначена ДО обходу header/footer, бо таблиці бувають і там (гриф
    # "ДЛЯ СЛУЖБОВОГО КОРИСТУВАННЯ" у клітинці поруч із номером примірника
    # -- реальна верстка офіційних листів), не лише в тілі документа.
    seen_ids, seen_refs = set(), []
    # Наскрізний лічильник таблиць документа: клітинка мусить знати СВОЮ
    # адресу (таблиця, рядок, стовпець), інакше багаторядкова таблиця
    # (книга обліку) нерозбірна за побудовою (R-A1-09: індекси були фізично
    # в руках і відкидались рядком blocks.append(text)).
    table_seq = [0]

    def walk_tables(tables):
        """Рекурсивно, бо таблиця може лежати ВСЕРЕДИНІ клітинки іншої
        таблиці (у бланках це звичайна річ), а `doc.tables` віддає лише
        верхній рівень -- вкладені раніше не читались узагалі й мовчки
        випадали з екстракції.

        Клітинка йде БЛОКОМ-СЛОВНИКОМ з адресою {table, row, col} -- увесь
        код нижче за течією вже вміє блоки-словники (шлях Surya/PDF), тож
        для екстракції це той самий текст, а адреса більше не викидається
        (R-A1-09). Абзаци лишаються рядками -- у них адреси немає.
        Об'єднана клітинка після дедуплікації несе адресу ПЕРШОГО входження
        (найлівіша колонка сітки під нею)."""
        for table in tables:
            t_index = table_seq[0]
            table_seq[0] += 1
            for r_index, row in enumerate(table.rows):
                for c_index, cell in enumerate(row.cells):
                    tc = cell._tc
                    if id(tc) in seen_ids:
                        continue
                    seen_ids.add(id(tc))
                    seen_refs.append(tc)
                    text = cell.text.strip()
                    if text:
                        blocks.append({"text": text, "bbox": None,
                                       "table": t_index, "row": r_index,
                                       "col": c_index})
                    if cell.tables:
                        walk_tables(cell.tables)

    # Кутовий штамп/шапка реального заповненого документа може бути
    # надрукована в header, не в тілі -- python-docx не читає header/footer
    # разом з doc.paragraphs, це треба явно. is_linked_to_previous
    # перевіряємо, бо секція без власного header успадковує header
    # попередньої секції -- без цієї перевірки один і той самий текст
    # задублювався б для кожної секції документа.
    # Дедуплікація за id() тут була б тим самим багом, що вже знайдений для
    # клітинок вище (lxml перевикористовує id звільнених проксі), тому
    # порівнюємо за самим ВМІСТОМ header/footer: він короткий, а різні секції
    # з однаковою шапкою -- це справді один і той самий текст.
    # first_page_header/footer -- ОКРЕМИЙ від звичайного header/footer вміст
    # (реальна властивість python-docx, перевірено), який Word показує лише
    # на 1-й сторінці секції, коли увімкнено "Особлива перша сторінка". Раніше
    # не читався взагалі -- офіційний лист із повною шапкою (адресат, гриф,
    # вихідний номер) саме на 1-й сторінці губив цей текст мовчки. Читаємо
    # лише коли прапорець дійсно увімкнений -- інакше python-docx все одно
    # віддає "порожній" проксі-об'єкт цієї секції, який Word не показує.
    seen_header_texts = set()
    header_footer_parts = []
    for section in doc.sections:
        header_footer_parts.append(section.header)
        header_footer_parts.append(section.footer)
        if section.different_first_page_header_footer:
            header_footer_parts.append(section.first_page_header)
            header_footer_parts.append(section.first_page_footer)

    for part in header_footer_parts:
        if part.is_linked_to_previous:
            continue
        part_texts = tuple(p.text.strip() for p in part.paragraphs if p.text.strip())
        if part_texts and part_texts not in seen_header_texts:
            seen_header_texts.add(part_texts)
            blocks.extend(part_texts)
        walk_tables(part.tables)

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            blocks.append(text)

    walk_tables(doc.tables)
    # join_block_texts, а не "\n".join: блоки-клітинки тепер словники.
    return join_block_texts(blocks), blocks


def sort_blocks_by_geometry(blocks):
    """v2: layout-aware впорядкування замість довіри до порядку списку.

    blocks: або list[str] (немає геометрії -- напр. docx, де порядок
    параграфів документа й так надійний, повертається без змін), або
    list[{"text": str, "bbox": (x1, y1, x2, y2)}] -- формат, який тепер
    повертає Surya для кожного блоку (bbox у координатах сторінки).

    Сортує за (y1, x1): зверху-вниз, зліва-направо за РЕАЛЬНОЮ позицією на
    сторінці, а не за тим, у якому порядку detection-модель внутрішньо
    згенерувала блоки. Дешева, детерміністична заміна повноцінних
    layout-моделей (LayoutLM/LiLT) -- без нової моделі й без GPU, лише
    геометрія, яку OCR вже рахує.

    Повертає list[{"text","bbox"}] (bbox БІЛЬШЕ НЕ відкидається -- раніше
    тут `[b["text"] for b in ordered]` губило геометрію одразу після
    сортування, і вона не доходила до extract_document взагалі; саме це
    не давало прив'язати значення до лейбла ГЕОМЕТРИЧНО, коли порядок
    блоків уводить в оману -- research-round-2026-08-12.md). list[str]
    повертається без змін, коли геометрії немає (docx)."""
    if not blocks or isinstance(blocks[0], str):
        return blocks
    missing = [b for b in blocks if not b.get("bbox")]
    if missing:
        # Явна помилка, а не тиха деградація до природного порядку: якщо
        # bbox зник (напр. інша версія/режим Surya), це має бути ПОМІЧЕНО,
        # бо без геометрії ми повертаємось до того самого класу багів
        # (сплутаний порядок блоків), який ця функція існує, щоб виправити.
        raise ValueError(
            f"{len(missing)} з {len(blocks)} блоків без bbox -- геометричне "
            "сортування неможливе. Перевірте версію/режим OCR-виклику."
        )
    return sorted(blocks, key=lambda b: (b["bbox"][1], b["bbox"][0]))


def _block_text(block) -> str:
    """block -- str або {"text","bbox"}; повертає сам текст незалежно від
    представлення. Потрібно, відколи sort_blocks_by_geometry перестала
    зрізати bbox перед поверненням.

    NUL (0x00) знімається ТУТ, на спільному виході текстів: це артефакт
    pypdf на окремих PDF, а не легальний символ документа. Заміряно
    24.08.2026: один NUL у НД ТЗІ 1.1-003-99 поклав запис у Postgres --
    text-поля NUL не приймають."""
    text = block["text"] if isinstance(block, dict) else block
    return (text or "").replace("\x00", "")


def join_block_texts(blocks) -> str:
    return "\n".join(_block_text(b) for b in blocks)


def file_sha256(path: str) -> str:
    """Ідентичність документа за вмістом, не за назвою файлу -- основа
    дедуплікації: ручний експорт людиною з АСКОД/Армія+ гарантує, що той
    самий документ рано чи пізно завантажать двічі, а подвійні факти
    подвоюють саме ті підрахунки, заради яких будується система."""
    digest = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def extract_pdf_blocks(path: str, ocr_fn=None, warnings=None, info=None,
                       max_ocr_pages=None):
    """PDF, рішення "текстовий шар чи скан" ПО КОЖНІЙ СТОРІНЦІ окремо:

    - сторінка з текстом -> `page.get_text("blocks")` дає блоки РАЗОМ з
      координатами, тому PDF отримує таку саму геометричну обробку, як OCR,
      без OCR;
    - сторінка без тексту (скан, вклеєне фото) -> рендериться в зображення й
      іде через ocr_fn.

    Раніше рішення приймалось один раз на весь документ за середнім числом
    символів, тому сторінка-скан усередині текстового PDF ніколи не
    потрапляла в OCR і мовчки випадала.

    Блоки сортуються ВСЕРЕДИНІ кожної сторінки, а сторінки склеюються за
    порядком: спільне сортування по всьому документу перемішало б сторінки
    між собою, бо y-координати на кожній сторінці починаються з нуля.

    max_ocr_pages -- СТЕЛЯ кількості сторінок, які підуть у розпізнавання
    (`ocr.max_pages` у конфізі; None = без межі, рівно попередня поведінка).
    Навіщо (рев'ю 22.08.2026, A-19): розпізнавання коштує хвилини НА СТОРІНКУ,
    а межі не було ніякої -- 200-сторінковий скан означав годинний прогін, у
    якому ще й немає посторінкового збереження: виняток будь-де в інжесті дає
    ОДИН `unresolved` на весь файл, тобто вся робота втрачається. Стеля --
    не оптимізація, а вихід: документ обробляється частково, і в попередженнях
    прямо сказано, ЯКІ сторінки не читались.

    PyMuPDF імпортується ліниво -- решта пайплайна працює без нього.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError as exc:
        raise ValueError(
            f"Для PDF потрібен PyMuPDF (pip install pymupdf): {path}"
        ) from exc

    import tempfile

    blocks = []
    pages_needing_ocr = []
    # Сторінки, у яких Є І текстовий шар, І растр, і сторінки зовсім порожні
    # (рев'ю 22.08.2026, C-05). Обидва класи раніше пропускались МОВЧКИ:
    # `continue` після текстової гілки навіть не питав `page.get_images()`.
    # Реальний вхід замовника саме такий -- сканований бланк, на який сканер
    # додав текстовий колонтитул, або PDF із текстової лицьової сторінки й
    # ВКЛЕЄНОЇ фотографії зворотного боку: 40 символів колонтитула досить, щоб
    # уся сторінка пішла текстовим шляхом, і зворотний бік зник без слідів.
    text_pages_with_images = []
    empty_pages = []
    pages_over_limit = []
    ocr_pages = 0
    with fitz.open(path) as doc:
        with tempfile.TemporaryDirectory() as tmpdir:
            for i, page in enumerate(doc):
                page_text = page.get_text() or ""
                if len(page_text.strip()) >= MIN_TEXT_CHARS_PER_PAGE:
                    if page.get_images():
                        text_pages_with_images.append(i + 1)
                    page_blocks = []
                    for b in page.get_text("blocks"):
                        # (x0, y0, x1, y1, text, block_no, block_type); 1 -- зображення
                        if len(b) >= 7 and b[6] != 0:
                            continue
                        text = (b[4] or "").strip()
                        if text:
                            # `page: i` -- кожна сторінка PDF рахує bbox з
                            # нуля (коментар нижче), тому геометричне
                            # порівняння МІЖ сторінками безглузде -- виміряний
                            # реальний баг: блок із СТОРІНКИ 2 "вирівнювався"
                            # з лейблом на СТОРІНЦІ 1 лише тому, що обидві
                            # сторінки рахують y з нуля (LEAVE-003.pdf,
                            # 2 сторінки). find_block_before_label/
                            # _geometric_candidate фільтрують за цим полем.
                            # `source` -- див. PDF_TEXT_SOURCE: блок цього
                            # походження не є абзацом, тож його межа не є
                            # межею поля.
                            page_blocks.append({"text": text, "bbox": (b[0], b[1], b[2], b[3]),
                                                "page": i, "source": PDF_TEXT_SOURCE})
                    blocks.extend(sort_blocks_by_geometry(page_blocks))
                    continue

                # Сторінка без текстового шару.
                # ПОРОЖНЯ сторінка (ні тексту, ні жодного зображення) -- НЕ
                # скан (R-B1-05): docx, збережений/перейменований у PDF, несе
                # порожню останню сторінку, і вона робила весь born-digital
                # документ «photo» (а photo -- це і чужа планка якості, і
                # вхід до черги «рукописне»). Скан за визначенням містить
                # растр; сторінка без растра -- просто порожній аркуш.
                if not page.get_images():
                    empty_pages.append(i + 1)
                    continue
                if ocr_fn is None:
                    pages_needing_ocr.append(i + 1)
                    continue
                if max_ocr_pages is not None and ocr_pages >= max_ocr_pages:
                    pages_over_limit.append(i + 1)
                    continue
                # ПОСТОРІНКОВИЙ ПРОГРЕС. Прогін на 200 сторінок без жодного
                # рядка виводу неможливо відрізнити від зависання -- саме це й
                # робило довгий скан «усе або нічого» ще й на вигляд.
                print(f"    OCR: сторінка {i + 1} з {doc.page_count}"
                      + (f" (стеля {max_ocr_pages})" if max_ocr_pages else ""),
                      flush=True)
                image_path = os.path.join(tmpdir, f"page_{i:03d}.png")
                page.get_pixmap(dpi=PDF_RENDER_DPI).save(image_path)
                ocr_blocks = ocr_fn(image_path)
                for block in ocr_blocks:
                    if isinstance(block, dict):
                        block["page"] = i
                blocks.extend(sort_blocks_by_geometry(ocr_blocks))
                ocr_pages += 1

    # Скільки сторінок пройшло через OCR -- визначає source_kind у БД-споживача
    # (CHECK electronic/photo): PDF з текстовим шаром це electronic, а
    # сканований PDF за суттю такий самий photo, як jpg зі смартфона, і
    # оцінювати його якість треба за тією самою планкою.
    # `scan_pages_detected` -- ОКРЕМО від ocr_pages: якщо `ocr.engine: none`,
    # скан-сторінки потрапляють у pages_needing_ocr, а НЕ проганяються через
    # OCR, тому ocr_pages лишається 0 і source_kind раніше хибно виходив
    # "electronic" для документа, що ЧАСТКОВО скан -- прямо суперечило
    # коментарю вище. Виявлення скан-сторінки не залежить від того, чи
    # налаштований OCR для її обробки.
    if info is not None:
        info["ocr_pages"] = ocr_pages
        info["scan_pages_detected"] = ocr_pages > 0 or bool(pages_needing_ocr)
        # Окремими ключами, а не в scan_pages_detected: сторінка з текстом і
        # растром -- це НЕ доказ скана. Замір по всіх 1122 сторінках PDF
        # репозиторію: 31 така сторінка, і всі -- у нормативних актах (штамп
        # або логотип на першій сторінці), у жодному факт-документі жодної.
        # Зарахувати їх у scan_pages_detected означало б перевести цілком
        # born-digital закон у `photo`, тобто в чужу планку якості й у чергу
        # «рукописне» -- це була б регресія, а не захист. Тому ключ окремий, а
        # рішення «чи OCR-ити растр на текстовій сторінці» лишається відкритим
        # (fixes-pipeline.md, «потребує рішення»).
        info["text_pages_with_images"] = list(text_pages_with_images)
        info["empty_pages"] = list(empty_pages)
        info["pages_over_ocr_limit"] = list(pages_over_limit)

    if pages_over_limit and warnings is not None:
        warnings.append(
            f"досягнуто стелі ocr.max_pages={max_ocr_pages}: не розпізнано "
            f"{len(pages_over_limit)} сторінок (№ "
            f"{', '.join(map(str, pages_over_limit[:10]))}"
            f"{' …' if len(pages_over_limit) > 10 else ''}) -- документ "
            "оброблено ЧАСТКОВО")

    if text_pages_with_images and warnings is not None:
        warnings.append(
            "растр без OCR на сторінках № "
            f"{', '.join(map(str, text_pages_with_images[:10]))}"
            f"{' …' if len(text_pages_with_images) > 10 else ''} "
            f"(усього {len(text_pages_with_images)}): на сторінці є І текстовий "
            "шар, І растр -- усе, що намальовано на растрі (вклеєна "
            "фотографія зворотного боку, штамп, підпис), у текст не потрапило")

    if pages_needing_ocr:
        if not blocks:
            raise ValueError(
                f"У PDF {path} немає текстового шару (скан), а OCR не налаштовано "
                "(ocr.engine: none у конфізі)."
            )
        # Частина сторінок -- скан, решта з текстом. Раніше умова була
        # `and not blocks`, тому за наявності хоч однієї текстової сторінки
        # скановані МОВЧКИ зникали: зворотний бік бланка (дата повернення,
        # зупинки) губився без жодного сліду. Тепер це попередження, яке
        # доходить до звіту й до запису.
        if warnings is not None:
            warnings.append(
                f"пропущено {len(pages_needing_ocr)} сторінок без текстового шару "
                f"(№ {', '.join(map(str, pages_needing_ocr))}) -- OCR не налаштовано")
    return join_block_texts(blocks), blocks


def load_document_blocks(path: str, ocr_fn=None, warnings=None, info=None,
                         max_ocr_pages=None):
    """ocr_fn: (image_path) -> list[{"text","bbox"}] або list[str];
    обов'язковий лише для зображень. Невідоме розширення -- явна помилка,
    без мовчазного fallback.

    Повертає (text, blocks) -- той самий формат, що очікує extract_document,
    незалежно від джерела. info (необов'язковий словник) наповнюється
    метаданими самого інжесту: ocr_pages -- скільки сторінок довелось
    розпізнавати, source_kind -- electronic/photo у термінах БД-споживача."""
    ext = os.path.splitext(path)[1].lower()

    if ext in DOCX_EXTS:
        if info is not None:
            info["ocr_pages"] = 0
            info["source_kind"] = "electronic"
        return extract_docx_blocks(path)

    if ext in PDF_EXTS:
        result = extract_pdf_blocks(path, ocr_fn=ocr_fn, warnings=warnings,
                                    info=info, max_ocr_pages=max_ocr_pages)
        if info is not None:
            # НЕ info.get("ocr_pages"): скан-сторінка, пропущена через
            # ocr.engine: none, ніколи не збільшує ocr_pages, але документ
            # усе одно частково скан -- source_kind має це відображати.
            info["source_kind"] = "photo" if info.get("scan_pages_detected") else "electronic"
        return result

    if ext in IMAGE_EXTS:
        if info is not None:
            info["ocr_pages"] = 1
            info["source_kind"] = "photo"
        if ocr_fn is None:
            raise ValueError(
                f"Файл {path} -- зображення, але OCR не налаштовано "
                "(ocr.engine: none у конфізі). Для зображень потрібен ocr.engine: surya."
            )
        # Сортування за геометрією централізоване тут, а не в ocr_fn, щоб
        # OCR-виклик лишався простим "зібрати блоки", не турбуючись про порядок.
        ordered_blocks = sort_blocks_by_geometry(ocr_fn(path))
        text = join_block_texts(ordered_blocks)
        # ВИХІД OCR ФІКСУЄТЬСЯ ЯК ФАКТ, а не перевіряється порогом. Причина
        # заміряна 13-14.08.2026: у прогоні 16 фото документи 10-16 пішли в
        # `unresolved` з нулем полів, а той самий файл в ОКРЕМОМУ процесі
        # розпізнався (991 символ, 25 блоків, бал ідентифікації 15 при порозі
        # 5). Тобто пакетний прогін деградує після ~дев'яти файлів -- і не
        # подає жодного сигналу: документ із порожнім розпізнаванням виглядає
        # рівно як документ невідомого типу.
        #
        # Кількість блоків і символів тепер видно в кожному записі, тож
        # деградацію можна ПОБАЧИТИ, не вгадуючи поріг. Порогові попередження
        # нижче -- додатковий шар, і вони саме попередження: статус не
        # змінюють, щоб не заміняти одну тиху поведінку іншою.
        if info is not None:
            info["ocr_blocks"] = len(ordered_blocks)
            info["ocr_chars"] = len(text)
        if warnings is not None:
            if not ordered_blocks:
                warnings.append(
                    "OCR не повернув ЖОДНОГО блоку -- це збій розпізнавання, "
                    "а не порожній документ")
            elif len(ordered_blocks) < MIN_PLAUSIBLE_OCR_BLOCKS or                     len(text) < MIN_PLAUSIBLE_OCR_CHARS:
                warnings.append(
                    f"OCR повернув підозріло мало: {len(ordered_blocks)} блоків, "
                    f"{len(text)} символів. Для аркуша бланка це малоймовірно -- "
                    "перевірте, чи не деградував OCR (див. known-weak-spots 2.18)")
        return text, ordered_blocks

    raise ValueError(f"Непідтримуване розширення файлу: {ext} ({path})")
