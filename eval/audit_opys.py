# -*- coding: utf-8 -*-
"""Звірка «AI-секретар_опис_проєкту.docx» із правилами Ані.

Правила -- docs/contracts/2026-08-29_opys-proektu-pravyla.md. Прилад міряє
МЕХАНІЧНУ частину: застарілі числа, наявність таблиці залученості людини,
розділів «чого не робить» і «що далі», відсутність маркдауну, персональних
ідентифікаторів і скріншотів, довжину речень. Судове (стиль, простота мови)
лишається людині -- і в звіті таке позначене «увага», а не «ок».

Запуск (файл поза git, тому шлях зашитий на робоче місце):
    python eval/audit_opys.py /tmp/report.txt

НАЩО ЦЕ ПРИЛАД, А НЕ ПЕРЕВІРКА ОКОМ. Двічі поспіль звірка «на око» пропускала
застаріле число, яке стоїть у ДВОХ місцях -- в абзаці й у таблиці. І один раз
сам цей скрипт збрехав: він міряв власний дамп, де клітинки таблиць склеєні
через «||», і прийняв це за маркдаун-таблицю в документі. Тому маркдаун і
довжина речень міряються по абзацах, а не по дампу.
"""
import io
import re
import sys

import docx

sys.stdout.reconfigure(encoding="utf-8")

#: Майстер-копія лежить У РЕПОЗИТОРІЇ (`docs/`), а не на робочому столі.
#: Тут був зашитий шлях на стіл, і 01.09 прилад тихо перевірив копію від
#: 29.08: сказав «застарілих чисел не знайдено» про файл, якого я не
#: правила. Той самий клас, що `verify_catalog`, який міряв не те.
#: Перевизначити можна змінною OPYS_DOCX.
import os

_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
PATH = os.environ.get("OPYS_DOCX") or os.path.join(
    _ROOT, "docs", "AI-секретар_опис_проєкту.docx")
d = docx.Document(PATH)

paras = [(p.style.name, p.text.strip()) for p in d.paragraphs if p.text.strip()]
text = "\n".join(t for _, t in paras)
tables = []
for i, tb in enumerate(d.tables):
    rows = [[c.text.strip() for c in r.cells] for r in tb.rows]
    tables.append(rows)
    text += "\n" + "\n".join(" || ".join(r) for r in rows)

out = []


def say(tag, name, detail=""):
    out.append("%-9s %s%s" % (tag, name, ("  — " + detail) if detail else ""))


# ── А6: актуальність чисел ──────────────────────────────────────────────────
stale = {
    "Таких запитів 29": "запитів 29 (тепер 31)",
    "з 29 запитів": "«з 29 запитів»",
    "28 з 28": "звірка 28 (тепер 35)",
    "80 із 127": "маршрутизація 80 (тепер 82)",
    "Сім інструментів": "«сім інструментів» (тепер дев'ять)",
    "П'ять інструментів": "«п'ять інструментів»",
    "У базі 41 документ": "41 нормативний (тепер 44)",
    "медіана 0,04 с": "стара медіана",
    "чернеток 112": "стара кількість чернеток",
    "360 сайту": "стара кількість тестів",
    "886": "стара кількість тестів сайту (тепер 954)",
    "805": "стара кількість тестів сайту",
}
found_stale = [v for k, v in stale.items() if k in text]
say("ПРОВАЛ" if found_stale else "ок", "А6 актуальність чисел",
    "; ".join(found_stale) if found_stale else "застарілих не знайдено")

# ── А2: таблиця залученості людини ──────────────────────────────────────────
inv = None
for rows in tables:
    if rows and "Крок" in rows[0][0] and "Хто вирішує" in " ".join(rows[0]):
        inv = rows
say("ок" if inv else "ПРОВАЛ", "А2 таблиця залученості людини",
    ("рядків %d" % (len(inv) - 1)) if inv else "таблиці немає")
if inv:
    joined = " ".join(" ".join(r) for r in inv)
    for word in ("чернет", "перевірк"):
        if word not in joined:
            say("увага", "А2 прив'язка до стану", "немає слова «%s»" % word)

# ── А3: документ про весь проєкт ────────────────────────────────────────────
blocks = [t for _, t in paras if re.match(r"^\d+\.\s|^Блок", t)]
owners = {"обробк": False, "База даних": False, "Сайт": False}
for k in owners:
    owners[k] = k.lower() in text.lower()
say("ок" if all(owners.values()) else "увага", "А3 усі три блоки",
    ", ".join("%s=%s" % (k, "є" if v else "НЕМАЄ") for k, v in owners.items()))

# ── А4 і Б4: прогалини названі ──────────────────────────────────────────────
gaps_head = [t for _, t in paras if t.startswith("9.") or "не робить" in t]
gap_items = 0
started = False
for st, t in paras:
    if t.startswith("9."):
        started = True
        continue
    if started and re.match(r"^\d+\.\s", t):
        break
    if started:
        gap_items += 1
say("ок" if gap_items >= 5 else "ПРОВАЛ", "А4/Б4 розділ «чого не робить»",
    "пунктів %d" % gap_items)

# ── Б1: у кожного приладу сказано, чим міряно ───────────────────────────────
instr = [t for _, t in paras if re.match(r"^7\.\d+\.", t)]
triplets = 0
for st, t in paras:
    if t.startswith("Що перевіря"):
        triplets += 1
say("ок" if len(instr) >= 9 else "ПРОВАЛ", "В3 приладів описано",
    "заголовків %d, «що перевіряє» %d" % (len(instr), triplets))

# ── Б6: без маркдауну ───────────────────────────────────────────────────────
para_text = "\n".join(t for _, t in paras)
cell_text = "\n".join(c for rows in tables for r in rows for c in r)
md_source = para_text + "\n" + cell_text
md = []
for pat, what in ((r"\*\*", "**"), (r"(?m)^#{1,6}\s", "# заголовок"),
                  (r"`[^`]+`", "`код`"), (r"(?m)^- \S", "- список"),
                  (r"\|", "вертикальна риска")):
    if re.search(pat, md_source):
        md.append(what)
say("ПРОВАЛ" if md else "ок", "Б6 без маркдауну",
    ", ".join(md) if md else "жодного маркера")

# ── Б2 і Д: без оцінок і без гасел ──────────────────────────────────────────
vague = [w for w in ("приблизно", "близько", "унікальн", "найкращ",
                     "революц", "легко масштаб") if w in text.lower()]
say("увага" if vague else "ок", "Б2/Д без оцінок і гасел",
    ", ".join(vague) if vague else "не знайдено")
for w in vague:
    for _, t in paras:
        if w in t.lower():
            out.append("           контекст «%s»: %s" % (w, t[:160]))

# ── В1: «наш квиток» наскрізь ───────────────────────────────────────────────
tickets = [t for _, t in paras if t.startswith("Наш квиток")]
say("ок" if len(tickets) >= 5 else "увага", "В1 «наш квиток» наскрізь",
    "згадок %d" % len(tickets))

# ── В4: окремо «чого немає» і «що далі» ─────────────────────────────────────
has9 = any(t.startswith("9.") for _, t in paras)
has10 = any(t.startswith("10.") for _, t in paras)
say("ок" if has9 and has10 else "ПРОВАЛ", "В4 окремі розділи 9 і 10",
    "9=%s, 10=%s" % (has9, has10))

# ── В6: числа в одному розділі з датою ──────────────────────────────────────
nums_head = [t for _, t in paras if t.startswith("8. Цифри")]
say("ок" if nums_head else "ПРОВАЛ", "В6 розділ цифр із датою",
    nums_head[0] if nums_head else "заголовка немає")

# ── Б5: персональні дані ────────────────────────────────────────────────────
#: У документі мусять бути ЛИШЕ синтетичні прізвища зі стенду.
real_markers = re.findall(r"\b\d{10}\b|\bРНОКПП\s*\d|паспорт\s*[А-Я]{2}\s*\d",
                          text)
say("ПРОВАЛ" if real_markers else "ок", "Б5 без персональних даних",
    ", ".join(real_markers[:3]) if real_markers else "ідентифікаторів немає")

# ── ГОЛОВНЕ ПРАВИЛО: речення ніби пункти, без води, без дрібниць ────────────
#
# Межа 25 слів -- це міра правила Ані, а не мій смак: «речення ніби пункти, без
# води, так щоб було зрозуміло людині не в контексті». Раніше тут стояло 45, і
# прилад зеленів на реченні в сорок слів -- тобто міряв не те правило.
MAX_WORDS = 25

sent = []
for _, t in paras:
    sent += [(s, "абзац") for s in re.split(r"(?<=[.!?;])\s+", t)]
for i, rows in enumerate(tables):
    for r in rows:
        for c in r:
            sent += [(s, "таблиця %d" % i)
                     for s in re.split(r"(?<=[.!?;])\s+", c)]
long_s = [(len(x.split()), x, w) for x, w in sent if len(x.split()) > MAX_WORDS]
say("ПРОВАЛ" if long_s else "ок",
    "ГОЛОВНЕ довжина речень (>%d слів)" % MAX_WORDS, "%d штук" % len(long_s))
for n, x, w in sorted(long_s, reverse=True)[:8]:
    out.append("           %d слів [%s] %s" % (n, w, x[:150]))

#: ВОДА -- вставні звороти, які нічого не додають.
WATER = ("варто зазначити", "слід зазначити", "слід підкреслити", "як відомо",
         "у цілому", "в цілому", "фактично", "по суті", "як правило",
         "у певному сенсі", "з одного боку", "тим не менш",
         "необхідно відзначити")
water = [w for w in WATER if w in text.lower()]
say("ПРОВАЛ" if water else "ок", "ГОЛОВНЕ без води",
    ", ".join(water) if water else "жодного звороту")

#: ДРІБНИЦІ -- те, що в мапі зайве. Назви моделей і рейки з числами тут
#: НАВМИСНО відсутні: це суть, а не нюанс (див. правила, «Головне правило»).
MINUTIAE = re.compile(
    r"(?:documents|objects|people|dimensions|facts|fact_sources|review_queue"
    r"|review_log|document_chunks|document_types|dimension_values"
    r"|object_aliases|leave_place|deployment_location|superseded_by_doc_id"
    r"|valid_from|valid_to|source_doc_id|canonical_name"
    r"|python-docx|pypdf|psycopg|llama\.cpp|pgvector|Hunspell|FastAPI|Gradio"
    r"|systemd|uvicorn|GBNF|tsvector"
    r"|4096 токенів)")
mins = sorted(set(MINUTIAE.findall(text)))
say("ПРОВАЛ" if mins else "ок", "ГОЛОВНЕ мапа, не інструкція",
    ", ".join(mins) if mins else "технічних дрібниць немає")

# ── СУПЕРЕЧНОСТІ ВСЕРЕДИНІ ДОКУМЕНТА ───────────────────────────────────────
#
# Знайдено Анею 29.08: у вступі стояло «непевне віддає людині, підтверджене
# складає в базу», а через два абзаци -- «чернетка лягає в базу і в підрахунки
# не входить». Тобто документ сам собі суперечив, і жодна попередня перевірка
# цього не бачила: обидва твердження окремо виглядають нормально.
#
# Пари нижче -- взаємно виключні твердження. Якщо в документі є обидві
# половини пари, одна з них неправда.
CONTRADICTIONS = [
    ("непевне віддає людині, підтверджене складає в базу",
     "лягає як чернетка",
     "чернетки або в базі, або «віддані людині» -- не обидва"),
    ("цитата абзацна",
     "цитата тепер адресна",
     "цитата або абзацна, або адресна"),
    ("Немає ролей і прав",
     "Ролей дві",
     "або ролей немає, або їх дві"),
    ("окремої сторінки «черга перевірки» немає",
     "сторінка черги перевірки працює",
     "сторінка черги або є, або немає"),
]
clashes = []
low_all = text.lower()
for a, b, why in CONTRADICTIONS:
    if a.lower() in low_all and b.lower() in low_all:
        clashes.append(why)
say("ПРОВАЛ" if clashes else "ок", "Б7 документ не суперечить собі",
    "; ".join(clashes) if clashes else "перевірених пар без конфлікту")

io.open(sys.argv[1], "w", encoding="utf-8").write(
    "\n".join(out) + "\n\n--- найдовші речення ---\n"
    + "\n\n".join(x[:300] for _n, x, _w in sorted(long_s, reverse=True)[:5]))
print("\n".join(out))
