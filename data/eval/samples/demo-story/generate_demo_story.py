# -*- coding: utf-8 -*-
"""
Генератор демо-набору до презентації: логічно зв'язана історія «від А до Я»
плюс масовка для обсягу бази.

Що робить (усе синтетичне, усі особи -- зі штатки `db/seeds/unit_roster.csv`):

  story/         17 .docx -- зв'язана історія (16 документів сценарію
                 + 1 навмисний виняток: особа, якої в штатці НЕМАЄ)
  story-pdf/     5 тих самих документів у .pdf (текстовий шар)
  story-photo/   4 тих самих документи як «знімок телефоном» (навмисно погана
                 якість: поворот, нерівне освітлення, розмиття, JPEG-артефакти)
  bulk/          ~130 .docx масовки (інші особи штатки, ~12% із порожнім
                 критичним полем)
  live/          3 документи резерву (docx + pdf + фото), які в базу НЕ
                 заливаються -- щоб на демо провести документ живцем
  еталони        data/eval/demo-story/per-document/<ID>.json -- той самий
                 формат, що data/eval/synthetic-2026-05/per-document/

Бланки НЕ вигадуються: docx збирається заливкою значень у порожні бланки
  data/eval/samples/leave/відпускний_шаблон.docx        (Додаток 30)
  data/eval/samples/deployment/посвідчення_відрядження.docx (Додаток 28)
у ті самі комірки, що й наявний корпус synthetic-2026-05 (звірено проти
LEAVE-001/011/016 і TRIP-001/014). Через це пайплайн упізнає ці документи
тими самими схемами, і демо показує наскрізний шлях, а не новий бланк.

PDF складається з ТОГО САМОГО набору рядків, у тому самому порядку, у якому
їх бачить `extract_docx_blocks` (значення -- окремим блоком ПЕРЕД друкованим
лейблом). Тому `block_before_label` працює на обох форматах однаково.

Дати -- ВІДНОСНО «сьогодні» (`--today`, дефолт -- системна дата): набір не
старіє, якщо демо перенесуть. Періоди підібрані так, щоб на день демо
хтось був у відпустці, хтось у відрядженні, хтось уже повернувся, десь був
перетин, а десь скасування вже подіяло.

Запуск (з кореня репозиторію):
    python data/eval/samples/demo-story/generate_demo_story.py --today 2026-08-28
    python data/eval/samples/demo-story/generate_demo_story.py --only story,pdf
"""
import argparse
import csv
import datetime as dt
import json
import os
import random

import numpy as np
from PIL import Image, ImageFilter
from docx import Document
from docx.shared import Pt

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.abspath(os.path.join(HERE, "..", "..", "..", ".."))

LEAVE_BLANK = os.path.join(ROOT, "data", "eval", "samples", "leave",
                           "відпускний_шаблон.docx")
DEPL_BLANK = os.path.join(ROOT, "data", "eval", "samples", "deployment",
                          "посвідчення_відрядження.docx")
ROSTER_CSV = os.path.join(ROOT, "db", "seeds", "unit_roster.csv")
EXPECTED_DIR = os.path.join(ROOT, "data", "eval", "demo-story", "per-document")

DIR_STORY = os.path.join(HERE, "story")
DIR_PDF = os.path.join(HERE, "story-pdf")
DIR_PHOTO = os.path.join(HERE, "story-photo")
DIR_BULK = os.path.join(HERE, "bulk")
DIR_LIVE = os.path.join(HERE, "live")

# Наша частина й підписант. Частина -- вигадана (літера + 4 цифри), підписант --
# зі штатки (командир батальйону), бо на бланку він друкується як
# «Власне ім'я ПРІЗВИЩЕ».
OUR_UNIT = "Ж3085"
SIGNER_RANK = "майор"
SIGNER_NAME = "Адам ДУТКА"

MONTHS_GEN = {
    1: "січня", 2: "лютого", 3: "березня", 4: "квітня", 5: "травня",
    6: "червня", 7: "липня", 8: "серпня", 9: "вересня", 10: "жовтня",
    11: "листопада", 12: "грудня",
}

# Пропис 1..31: рівно те, що вміє pipeline/normalization (UKR_NUMBER_WORDS +
# складені числа). Апостроф -- типографський ’, як у документах.
_ONES = {1: "один", 2: "два", 3: "три", 4: "чотири", 5: "п’ять", 6: "шість",
         7: "сім", 8: "вісім", 9: "дев’ять", 10: "десять", 11: "одинадцять",
         12: "дванадцять", 13: "тринадцять", 14: "чотирнадцять",
         15: "п’ятнадцять", 16: "шістнадцять", 17: "сімнадцять",
         18: "вісімнадцять", 19: "дев’ятнадцять", 20: "двадцять"}


def days_in_words(number):
    if number in _ONES:
        return _ONES[number]
    if 21 <= number <= 29:
        return "двадцять " + _ONES[number - 20]
    if number == 30:
        return "тридцять"
    if number == 31:
        return "тридцять один"
    raise ValueError(f"пропис не передбачений для {number} (набір -- дні, 1..31)")


# ---------------------------------------------------------------------------
# Штатка
# ---------------------------------------------------------------------------
def load_roster():
    with open(ROSTER_CSV, encoding="utf-8") as fh:
        rows = list(csv.DictReader(fh))
    return {r["service_id"]: r for r in rows}


def person_slots(person, *, in_roster=True):
    """Словник, з якого складаються надруковані рядки бланка для цієї особи."""
    surname = person["last_name"]
    given = person["first_name"]
    patronymic = person["patronymic"]
    rank = person["rank"]
    female = person["gender"] == "жіноча"
    return {
        "service_id": person.get("service_id"),
        "rank": rank,
        "surname": surname,
        "given": given,
        "patronymic": patronymic,
        "gender": person["gender"],
        "position": person["position_title"],
        "subdivision": person.get("subdivision", ""),
        "in_roster": in_roster,
        "PERSON_FULL": f"{rank} {surname.upper()} {given} {patronymic}",
        "PERSON_SHORT": f"{rank} {surname.upper()} {given[0]}.{patronymic[0]}.",
        "RELEASED": "звільнена" if female else "звільнений",
        "OBLIGED": "зобов’язана" if female else "зобов’язаний",
        "SENT_TO": "відрядженій" if female else "відрядженому",
        "full_name": f"{surname} {given} {patronymic}",
    }


# ---------------------------------------------------------------------------
# Заливка docx
# ---------------------------------------------------------------------------
def _write_par(par, text):
    """Ставить текст у наявний абзац, зберігаючи його шрифт (перший run)."""
    runs = par.runs
    if runs:
        runs[0].text = text
        for extra in runs[1:]:
            extra.text = ""
    else:
        run = par.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def _set_par(doc, index, text):
    _write_par(doc.paragraphs[index], text)


def _set_cell(table, row, col, text):
    _write_par(table.rows[row].cells[col].paragraphs[0], text)


def _date_slots(prefix, date):
    """Три слоти бланка на одну дату: день / місяць прописом / дві цифри року."""
    if date is None:
        return {f"{prefix}_D": "", f"{prefix}_M": "", f"{prefix}_Y": ""}
    return {f"{prefix}_D": f"{date.day:02d}",
            f"{prefix}_M": MONTHS_GEN[date.month],
            f"{prefix}_Y": f"{date.year % 100:02d}"}


def _range_line(v, a, b):
    """Рядок діапазону бланка. Порожні слоти дають рівно те саме, що корпус:
    'з “”  20 р.  по “”  20 р.' -- скелет порожнього слота, який схема
    впізнає своїм empty_pattern."""
    return (f"з “{v[a + '_D']}” {v[a + '_M']} 20{v[a + '_Y']} р.  "
            f"по “{v[b + '_D']}” {v[b + '_M']} 20{v[b + '_Y']} р.")


def fill_leave_docx(path, v):
    doc = Document(LEAVE_BLANK)
    signer = f"{v['SIGNER_RANK']}                                        {v['SIGNER_NAME']}"
    _set_par(doc, 6, f"військова частина {v['UNIT']}")
    _set_par(doc, 9, f"№ {v['DOC_NUMBER']}    від {v['ISSUE_DATE']}")
    _set_par(doc, 16, signer)
    _set_par(doc, 31, signer)

    t = doc.tables[0]
    _set_cell(t, 0, 0, v["PERSON_FULL"])
    _set_cell(t, 2, 0, v["RELEASED"])
    _set_cell(t, 2, 1, v["LEAVE_TYPE"])
    _set_cell(t, 4, 0, v["LEAVE_PLACE"])
    _set_cell(t, 6, 2, v["DAYS_WORDS"])
    _set_cell(t, 8, 0, _range_line(v, "START", "END"))
    _set_cell(t, 9, 4, v["PERSON_SHORT"])
    _set_cell(t, 11, 0, f"{v['OBLIGED']} прибути до місця служби у ")
    _set_cell(t, 13, 0, v["RETURN_UNIT"])
    _set_cell(t, 15, 0, f"“{v['RET_D']}” {v['RET_M']} 20{v['RET_Y']} р.")
    _set_cell(t, 17, 5, v["VPD"])
    _set_cell(doc.tables[1], 0, 1, v["COMPANIONS"])
    doc.save(path)


def fill_deployment_docx(path, v):
    doc = Document(DEPL_BLANK)
    signer = f"{v['SIGNER_RANK']}                                        {v['SIGNER_NAME']}"
    _set_par(doc, 7, f"військова частина {v['UNIT']}")
    _set_par(doc, 11, f"№ {v['DOC_NUMBER']}    від {v['ISSUE_DATE']}")
    _set_par(doc, 16, signer)

    t = doc.tables[0]
    _set_cell(t, 0, 1, v["PERSON_FULL"])
    _set_cell(t, 2, 0, v["POSITION"])
    _set_cell(t, 4, 0, f"{v['SENT_TO']} до ")
    _set_cell(t, 4, 2, v["DEST"])
    _set_cell(t, 6, 0, v["DEST_ORG"])
    _set_cell(t, 8, 0,
              f"Термін відрядження “{v['DAYS']}” днів   "
              f"з “{v['START_D']}” {v['START_M']} 20{v['START_Y']} р. "
              f"по “{v['END_D']}” {v['END_M']} 20{v['END_Y']} р.")
    _set_cell(t, 9, 0, v["PURPOSE"])
    _set_cell(t, 12, 0, f"Підстава відрядження: {v['ORDER_BASIS']}")
    doc.save(path)


# ---------------------------------------------------------------------------
# Ті самі документи у PDF (текстовий шар)
# ---------------------------------------------------------------------------
# Порядок рядків = порядок блоків, у якому їх бачить пайплайн на docx:
# значення стоїть блоком ПЕРЕД своїм друкованим лейблом. Порожній слот рядка
# не дає взагалі (у docx порожня комірка блоком не стає).
def _lines(*items):
    return [x for x in items if x is not None and x != ""]


# Сторінка PDF -- список ГРУП, група -- один блок тексту (один виклик
# insert_htmlbox = один блок на витягу). Значення й його друкований лейбл
# стоять В ОДНІЙ групі НАВМИСНО: так само їх складає Word у наявному корпусі
# (перевірено на LEAVE-001.pdf і TRIP-001.pdf), і саме на цьому працює
# основний шлях `block_before_label` -- «лейбл не перший рядок свого блоку».
# Коли лейбл лишався ОКРЕМИМ блоком, `_extend_across_pdf_wrap` дотягував
# значення назад через межу блоку й приклеював до нього попередній рядок
# бланка ("зобов’язана прибути до місця служби у військова частина Ж3085"):
# у facts.value для тієї самої частини їхало два різні написання залежно від
# формату файлу.
def leave_pdf_pages(v):
    signer = f"{v['SIGNER_RANK']}                    {v['SIGNER_NAME']}"
    page1 = [
        _lines("Додаток 30",
               "до Інструкції з діловодства у Збройних Силах України",
               "(підпункт 2.8.9)"),
        _lines("Кутовий штамп", "військової частини (установи)",
               f"військова частина {v['UNIT']}"),
        _lines("Відпускний квиток"),
        _lines(f"№ {v['DOC_NUMBER']}    від {v['ISSUE_DATE']}"),
        _lines("Дійсний у разі пред’явлення документа, який засвідчує особу."),
        _lines(v["PERSON_FULL"],
               "(військове звання, прізвище, ім’я та по батькові)"),
        _lines(v["RELEASED"], v["LEAVE_TYPE"],
               "(вид відпустки та найменування населеного пункту,"),
        _lines(v["LEAVE_PLACE"], "до якого звільнено військовослужбовця)"),
        _lines("терміном на ", v["DAYS_WORDS"], "(кількість днів прописом)"),
        _lines(_range_line(v, "START", "END")),
        _lines("Після закінчення строку відпустки", v["PERSON_SHORT"]),
        _lines("(військове звання, прізвище та ініціали ",
               f"{v['OBLIGED']} прибути до місця служби у "),
        _lines(v["RETURN_UNIT"],
               "(найменування військової частини або населеного пункту)"),
        _lines(f"“{v['RET_D']}” {v['RET_M']} 20{v['RET_Y']} р.",
               "(дата повернення)"),
        _lines("Для проїзду видано військові перевізні документи за №",
               v["VPD"]),
        _lines("Разом з ", v["COMPANIONS"],
               "(військове звання, прізвище та ініціали", "прямують"),
        _lines("Командир (начальник)", signer,
               "(військове звання підпис Власне ім’я  ПРІЗВИЩЕ)", "М.П."),
    ]
    page2 = [
        _lines("2", "Продовження додатка 30",
               "Зворотний бік відпускного квитка"),
        _lines("Командир (начальник)", signer,
               "(військове звання підпис Власне ім’я  ПРІЗВИЩЕ)", "М.П."),
        _lines("Відмітка про постановку на облік та зняття з обліку"),
    ]
    return [[g for g in page1 if g], [g for g in page2 if g]]


def deployment_pdf_pages(v):
    signer = f"{v['SIGNER_RANK']}                    {v['SIGNER_NAME']}"
    page1 = [
        _lines("Додаток 28",
               "до Інструкції з діловодства у Збройних Силах України",
               "(пункт 2.8.8.3)"),
        _lines("Кутовий штамп ", "військової частини (установи)",
               f"військова частина {v['UNIT']}"),
        _lines("Посвідчення про відрядження"),
        _lines(f"№ {v['DOC_NUMBER']}    від {v['ISSUE_DATE']}"),
        _lines("Дійсно в разі пред’явлення документа, який засвідчує особу."),
        _lines("Видано", v["PERSON_FULL"],
               "(військове звання, прізвище ім’я по батькові)"),
        _lines(v["POSITION"], "(посада, місце роботи)"),
        _lines(f"{v['SENT_TO']} до ", v["DEST"], "(пункти призначень)"),
        _lines(v["DEST_ORG"],
               "(найменування військової частини, установи, організації)"),
        _lines(f"Термін відрядження “{v['DAYS']}” днів   "
               f"з “{v['START_D']}” {v['START_M']} 20{v['START_Y']} р. "
               f"по “{v['END_D']}” {v['END_M']} 20{v['END_Y']} р."),
        _lines(v["PURPOSE"], "(мета відрядження)"),
        _lines(f"Підстава відрядження: {v['ORDER_BASIS']}"),
        _lines("Командир (начальник)", signer,
               "(військове звання підпис Власне ім’я  ПРІЗВИЩЕ)", "М.П."),
    ]
    page2 = [
        _lines("2", "Продовження додатка 28",
               "Зворотний бік посвідчення про відрядження"),
        _lines("Примітка: у разі виїзду в декілька пунктів відмітки про "
               "прибуття та вибуття робляться окремо в кожному з них."),
    ]
    return [[g for g in page1 if g], [g for g in page2 if g]]


def render_pdf(path, pages, *, fontsize=11, line_h=15, gap=11):
    """Кладе групи рядків на A4: ОДНА група = один виклик = один блок на витягу.

    `insert_htmlbox` із вбудованим шрифтом mupdf (`font-family: serif`), а НЕ
    `insert_text` з times.ttf: при вбудованому TTF PyMuPDF віддає на витягу
    пробіл як NBSP (\\xa0), а схеми шукають літерали зі ЗВИЧАЙНИМ пробілом
    ("(кількість днів", "документи за №") -- на таких PDF тихо губились
    `duration_days` і `travel_document_number`, які в docx того самого
    документа є.
    """
    import fitz  # локально: PyMuPDF потрібен лише для pdf/фото
    import html

    doc = fitz.open()
    for groups in pages:
        page = doc.new_page(width=595, height=842)
        y = 50
        for group in groups:
            height = len(group) * line_h + 6
            body = "<br>".join(html.escape(line) for line in group)
            page.insert_htmlbox(
                fitz.Rect(56, y, 545, y + height),
                f'<div style="font-family:serif;font-size:{fontsize}px;'
                f'line-height:1.25">{body}</div>')
            y += height + gap
    # subset_fonts + garbage/deflate: інакше кожен PDF важить ~1.3 МБ (повний
    # вбудований шрифт), а в git їх шість.
    try:
        doc.subset_fonts()
    except Exception:
        pass
    doc.save(path, garbage=4, deflate=True, clean=True)
    doc.close()


# ---------------------------------------------------------------------------
# «Знімок телефоном»: навмисно погана якість
# ---------------------------------------------------------------------------
# 150 dpi, а не 200: і ближче до знімка аркуша телефоном, і файл утричі менший.
def pdf_first_page_image(pdf_path, dpi=150):
    import fitz

    doc = fitz.open(pdf_path)
    page = doc[0]
    pix = page.get_pixmap(dpi=dpi)
    img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
    doc.close()
    return img


def degrade_photo(img, rng, *, angle, blur, shadow, noise, crop_px):
    """Псує рендер так, як псує знімок телефоном: поворот, нерівне освітлення
    з тінню на частині аркуша, розмиття, шум сенсора, обрізаний край."""
    img = img.rotate(angle, resample=Image.BICUBIC, expand=True,
                     fillcolor=(232, 230, 226))

    arr = np.asarray(img).astype(np.float32)
    height, width = arr.shape[:2]

    # Нерівне освітлення: яскраво в одному куті, темніше в протилежному.
    ys = np.linspace(0.0, 1.0, height)[:, None]
    xs = np.linspace(0.0, 1.0, width)[None, :]
    light = 1.06 - 0.30 * ys - 0.16 * xs

    # Тінь смугою (рука/край столу над аркушем) -- м'який перехід.
    band_center = rng.uniform(0.45, 0.75)
    band_width = rng.uniform(0.10, 0.22)
    band = np.exp(-((ys - band_center) ** 2) / (2 * band_width ** 2))
    light = light - shadow * band

    arr = np.clip(arr * light[:, :, None], 0, 255)

    img = Image.fromarray(arr.astype(np.uint8))
    img = img.filter(ImageFilter.GaussianBlur(blur))

    arr = np.asarray(img).astype(np.float32)
    rs = np.random.RandomState(rng.randrange(2 ** 31))
    arr += rs.normal(0.0, noise, arr.shape)
    img = Image.fromarray(np.clip(arr, 0, 255).astype(np.uint8))

    # Обрізаний край: телефон майже завжди зрізає частину аркуша.
    w, h = img.size
    img = img.crop((crop_px, crop_px // 2, w - crop_px // 3, h - crop_px // 4))
    return img


PHOTO_PROFILES = {
    # (кут, розмиття, тінь, шум, обрізка, формат, якість)
    "normal": dict(angle=1.6, blur=1.0, shadow=0.20, noise=4.0, crop_px=26,
                   fmt="jpg", quality=62),
    "normal2": dict(angle=-2.4, blur=1.3, shadow=0.26, noise=5.0, crop_px=34,
                    fmt="jpg", quality=55),
    "png": dict(angle=2.2, blur=1.1, shadow=0.22, noise=4.5, crop_px=30,
                fmt="png", quality=None),
    # Навмисно ГІРШИЙ за решту -- цікаво, де межа OCR.
    "worst": dict(angle=-3.0, blur=2.6, shadow=0.38, noise=9.0, crop_px=48,
                  fmt="jpg", quality=32),
}


def make_photo(pdf_path, out_dir, doc_id, profile_name, rng):
    profile = dict(PHOTO_PROFILES[profile_name])
    fmt = profile.pop("fmt")
    quality = profile.pop("quality")
    img = degrade_photo(pdf_first_page_image(pdf_path), rng, **profile)
    out = os.path.join(out_dir, f"{doc_id}.{fmt}")
    if fmt == "jpg":
        img.save(out, "JPEG", quality=quality, subsampling=2)
    else:
        img.save(out, "PNG")
    return out


# ---------------------------------------------------------------------------
# Складання значень одного документа
# ---------------------------------------------------------------------------
def build_leave_values(spec, who, today):
    start = today + dt.timedelta(days=spec["start"]) if spec["start"] is not None else None
    end = today + dt.timedelta(days=spec["end"]) if spec["end"] is not None else None
    issue = today + dt.timedelta(days=spec["issue"])
    ret = today + dt.timedelta(days=spec["ret"]) if spec.get("ret") is not None else None
    days = (end - start).days + 1 if (start and end) else None

    v = {
        "UNIT": OUR_UNIT,
        "DOC_NUMBER": spec["number"],
        "ISSUE_DATE": issue.strftime("%d.%m.%Y"),
        "PERSON_FULL": who["PERSON_FULL"],
        "PERSON_SHORT": who["PERSON_SHORT"],
        "RELEASED": who["RELEASED"],
        "OBLIGED": who["OBLIGED"],
        "LEAVE_TYPE": spec["leave_type"],
        "LEAVE_PLACE": spec["place"],
        "DAYS_WORDS": days_in_words(days) if days else "",
        "RETURN_UNIT": f"військова частина {OUR_UNIT}",
        "VPD": spec.get("vpd", "не видавались"),
        "COMPANIONS": spec.get("companions", "—"),
        "SIGNER_RANK": SIGNER_RANK,
        "SIGNER_NAME": SIGNER_NAME,
    }
    v.update(_date_slots("START", start))
    v.update(_date_slots("END", end))
    v.update(_date_slots("RET", ret))

    # Навмисні прогалини: слот на бланку лишається порожнім.
    for gap in spec.get("gaps", ()):
        if gap == "place":
            v["LEAVE_PLACE"] = ""
        elif gap == "person":
            v["PERSON_FULL"] = v["PERSON_SHORT"] = ""
        elif gap == "dates":
            v.update(_date_slots("START", None))
            v.update(_date_slots("END", None))
            v["DAYS_WORDS"] = ""
        else:
            raise ValueError(f"невідома прогалина {gap}")

    truth = {
        "номер_документа": spec["number"],
        "дата_видачі": issue.isoformat(),
        "початок": start.isoformat() if start else "",
        "кінець": end.isoformat() if end else "",
        "повернення": ret.isoformat() if ret else "",
        "днів": days if days else "",
        "місце": v["LEAVE_PLACE"],
        "підстава": spec["leave_type"],
        "організація": "",
        "супутники": v["COMPANIONS"],
        "військова_частина": OUR_UNIT,
        "підписант": f"{SIGNER_RANK} {SIGNER_NAME}",
    }
    return v, truth


def build_deployment_values(spec, who, today):
    start = today + dt.timedelta(days=spec["start"]) if spec["start"] is not None else None
    end = today + dt.timedelta(days=spec["end"]) if spec["end"] is not None else None
    issue = today + dt.timedelta(days=spec["issue"])
    order = today + dt.timedelta(days=spec["order"])
    days = (end - start).days + 1 if (start and end) else None

    v = {
        "UNIT": OUR_UNIT,
        "DOC_NUMBER": spec["number"],
        "ISSUE_DATE": issue.strftime("%d.%m.%Y"),
        "PERSON_FULL": who["PERSON_FULL"],
        "SENT_TO": who["SENT_TO"],
        "POSITION": f"{who['position']}, військова частина {OUR_UNIT}",
        "DEST": spec["dest"],
        "DEST_CITY": spec["dest"],
        "DEST_ORG": spec["dest_org"],
        "DAYS": str(days) if days else "",
        "PURPOSE": spec["purpose"],
        "ORDER_BASIS": (f"наказ командира військової частини {OUR_UNIT} "
                        f"від {order.strftime('%d.%m.%Y')} № {spec['order_number']}"),
        "SIGNER_RANK": SIGNER_RANK,
        "SIGNER_NAME": SIGNER_NAME,
    }
    v.update(_date_slots("START", start))
    v.update(_date_slots("END", end))

    for gap in spec.get("gaps", ()):
        if gap == "dest":
            v["DEST"] = v["DEST_CITY"] = ""
        elif gap == "person":
            v["PERSON_FULL"] = ""
        elif gap == "dates":
            v.update(_date_slots("START", None))
            v.update(_date_slots("END", None))
            v["DAYS"] = ""
        else:
            raise ValueError(f"невідома прогалина {gap}")

    truth = {
        "номер_документа": spec["number"],
        "дата_видачі": issue.isoformat(),
        "початок": start.isoformat() if start else "",
        "кінець": end.isoformat() if end else "",
        "повернення": end.isoformat() if end else "",
        "днів": days if days else "",
        "місце": v["DEST"],
        "підстава": spec["purpose"],
        "організація": spec["dest_org"],
        "супутники": "—",
        "військова_частина": OUR_UNIT,
        "підписант": f"{SIGNER_RANK} {SIGNER_NAME}",
    }
    return v, truth


LEAVE_PRINTED_KEYS = [
    "UNIT", "DOC_NUMBER", "ISSUE_DATE", "PERSON_FULL", "PERSON_SHORT",
    "RELEASED", "OBLIGED", "LEAVE_TYPE", "LEAVE_PLACE", "DAYS_WORDS",
    "START_D", "START_M", "START_Y", "END_D", "END_M", "END_Y",
    "RET_D", "RET_M", "RET_Y", "RETURN_UNIT", "VPD", "COMPANIONS",
    "SIGNER_RANK", "SIGNER_NAME",
]
DEPL_PRINTED_KEYS = [
    "UNIT", "DOC_NUMBER", "ISSUE_DATE", "PERSON_FULL", "POSITION", "SENT_TO",
    "DEST", "DEST_ORG", "DAYS", "START_D", "START_M", "START_Y",
    "END_D", "END_M", "END_Y", "PURPOSE", "ORDER_BASIS", "DEST_CITY",
    "SIGNER_RANK", "SIGNER_NAME",
]


def expected_json(doc_id, spec, who, v, truth):
    kind = spec["kind"]
    keys = LEAVE_PRINTED_KEYS if kind == "leave" else DEPL_PRINTED_KEYS
    return {
        "id": doc_id,
        "тип": ("відпускний квиток" if kind == "leave"
                else "посвідчення про відрядження"),
        "бланк": "Додаток 30" if kind == "leave" else "Додаток 28",
        "категорія": spec.get("category", "правильний"),
        "вада": spec.get("defect"),
        "пара": spec.get("pair"),
        "чинний": spec.get("valid", True),
        "примітка": spec.get("note", ""),
        "людина": {
            "service_id": who["service_id"],
            "звання": who["rank"],
            "ПІБ": who["full_name"],
            "стать": who["gender"],
            "посада": who["position"],
            "підрозділ": who["subdivision"],
            "є_в_реєстрі": who["in_roster"],
        },
        "надруковано": {k: v.get(k, "") for k in keys},
        "правильні_відповіді": truth,
    }


# ---------------------------------------------------------------------------
# СЦЕНАРІЙ. Зсуви -- дні відносно «сьогодні».
# ---------------------------------------------------------------------------
# Особи -- дослівно зі штатки (звання + ПІБ + посада), інакше кожен документ
# створив би в базі НОВУ особу замість зіставлення з реєстром.
STORY_PEOPLE = {
    "havrysh": "UNIT-0001",     # сержант Гавриш Адам Станіславович
    "malyshko": "UNIT-0031",    # рядовий Малишко Камілла Омелянівна
    "pryimak_o": "UNIT-0179",   # молодший лейтенант Приймак Остап Русланович
    "pryimak_ye": "UNIT-0270",  # старший сержант Приймак Єлисей Романович
    "yurchuk": "UNIT-0054",     # старший сержант Юрчук Лариса Трохимівна
    "yaremkiv": "UNIT-0076",    # старший лейтенант Яремків Лариса Данівна
}

# Особа, якої в штатці НЕМА -- рівно один документ набору (черга рев'ю
# `new_person`). ПІБ вигаданий і в реєстрі не зустрічається.
OUTSIDER = {
    "service_id": None, "rank": "молодший сержант", "last_name": "Крижанівський",
    "first_name": "Тарас", "patronymic": "Богданович", "gender": "чоловіча",
    "position_title": "командир відділення", "subdivision": "",
}

STORY = [
    # --- Гавриш: відпустка -> повернувся -> відрядження (зараз у дорозі) ---
    dict(id="DEMO-01", kind="leave", who="havrysh", number="101",
         issue=-28, start=-25, end=-14, ret=-13,
         leave_type="щорічна основна відпустка за 2026 рік",
         place="м. Тихолісся", vpd="4180/26",
         note="Закрита відпустка: дата повернення вписана."),
    dict(id="DEMO-02", kind="deployment", who="havrysh", number="207",
         issue=-4, order=-5, order_number="391", start=-2, end=3,
         dest="м. Кривоярськ", dest_org="військова частина Т5140",
         purpose="отримання матеріально-технічних засобів",
         note="Триває на день демо -- «хто зараз у відрядженні»."),
    dict(id="DEMO-11", kind="leave", who="havrysh", number="136",
         issue=-1, start=10, end=16,
         leave_type="щорічна додаткова відпустка за 2026 рік",
         place="с. Соснова Гряда", vpd="4402/26",
         note="Третій документ тієї самої особи -- послідовність у часі."),

    # --- Малишко: відпустка, яку СКАСОВУЄ інший квиток ---
    dict(id="DEMO-03", kind="leave", who="malyshko", number="118",
         issue=-10, start=-1, end=12,
         leave_type="відпустка за сімейними обставинами",
         place="м. Кривоярськ", vpd="4211/26", companions="діти — 1",
         category="пара", valid=False,
         pair={"group": "Д1", "role": "скасований",
               "relation": "квиток анульований", "replaced_by": "DEMO-04"},
         note="Скасований квиток. У самому документі ознаки скасування НЕМА -- "
              "пару можна знайти лише запитом по всіх документах."),
    dict(id="DEMO-04", kind="leave", who="malyshko", number="131",
         issue=-3, start=-1, end=3,
         leave_type="відпустка за сімейними обставинами "
                    "(виданий замість анульованого квитка № 118)",
         place="м. Кривоярськ", vpd="4266/26", companions="діти — 1",
         category="пара",
         pair={"group": "Д1", "role": "чинний",
               "relation": "квиток анульований", "replaces": "DEMO-03"},
         note="Чинний квиток: строк КОРОТШИЙ за скасований."),
    dict(id="DEMO-15", kind="leave", who="malyshko", number="141",
         issue=2, start=17, end=23,
         leave_type="щорічна основна відпустка за 2026 рік",
         place="м. Кривоярськ", vpd="4455/26", companions="діти — 1"),

    # --- Двоє однофамільців із відпустками, що перетинаються ---
    dict(id="DEMO-05", kind="leave", who="pryimak_o", number="107",
         issue=-12, start=-8, end=5,
         leave_type="щорічна основна відпустка за 2026 рік",
         place="м. Заріччя-Долинське", vpd="не видавались",
         note="У відпустці на день демо; перетин із DEMO-06 у кінці строку."),
    dict(id="DEMO-06", kind="leave", who="pryimak_ye", number="112",
         issue=-2, start=1, end=14,
         leave_type="відпустка у зв’язку з навчанням",
         place="м. Сухобрід", vpd="4288/26",
         note="Однофамілець DEMO-05, інше по батькові -- два різні об'єкти."),
    dict(id="DEMO-16", kind="deployment", who="pryimak_o", number="226",
         issue=5, order=4, order_number="415", start=10, end=12,
         dest="м. Малий Ясенець", dest_org="навчальний центр (в/ч Т5140)",
         purpose="участь у зборах командирів взводів"),
    dict(id="DEMO-12", kind="deployment", who="pryimak_ye", number="231",
         issue=14, order=13, order_number="428", start=18, end=21,
         dest="м. Заріччя-Долинське", dest_org="військова частина К7719",
         purpose="супроводження вантажу"),

    # --- Юрчук: НАВМИСНА ПРОГАЛИНА (порожнє критичне поле «місце») ---
    dict(id="DEMO-07", kind="leave", who="yurchuk", number="109",
         issue=-7, start=-4, end=9,
         leave_type="відпустка за сімейними обставинами",
         place="м. Сухобрід", vpd="4230/26", gaps=("place",),
         category="зіпсований", defect="empty_fields",
         note="Навмисна прогалина: населений пункт у слот не вписаний. "
              "Критичне поле -> запис не підтверджений, у чаті «непідтверджених: 1»."),
    dict(id="DEMO-14", kind="deployment", who="yurchuk", number="228",
         issue=8, order=7, order_number="420", start=13, end=17,
         dest="м. Тихолісся", dest_org="військова частина Р2268",
         purpose="передача обладнання за актом"),

    # --- Яремків: відрядження, переоформлене іншим посвідченням ---
    dict(id="DEMO-08", kind="deployment", who="yaremkiv", number="201",
         issue=-18, order=-19, order_number="352", start=-16, end=-12,
         dest="м. Тихолісся",
         dest_org="Центральна база зберігання майна (в/ч Р2268)",
         purpose="отримання засобів індивідуального захисту",
         note="Завершене відрядження."),
    dict(id="DEMO-09", kind="deployment", who="yaremkiv", number="214",
         issue=-7, order=-8, order_number="377", start=3, end=7,
         dest="м. Сухобрід", dest_org="військова частина К7719",
         purpose="проходження курсу підвищення кваліфікації",
         category="пара", valid=False,
         pair={"group": "Д2", "role": "скасований",
               "relation": "відрядження переоформлене", "replaced_by": "DEMO-10"},
         note="Переоформлене посвідчення. Ознаки скасування в ньому НЕМА."),
    dict(id="DEMO-10", kind="deployment", who="yaremkiv", number="223",
         issue=-2, order=-3, order_number="402", start=4, end=6,
         dest="м. Сухобрід", dest_org="військова частина К7719",
         purpose="проходження курсу підвищення кваліфікації "
                 "(переоформлено замість посвідчення № 214)",
         category="пара",
         pair={"group": "Д2", "role": "чинний",
               "relation": "відрядження переоформлене", "replaces": "DEMO-09"},
         note="Чинне посвідчення: інші дати, ніж у скасованому."),
    dict(id="DEMO-13", kind="leave", who="yaremkiv", number="138",
         issue=3, start=8, end=21,
         leave_type="щорічна основна відпустка за 2026 рік",
         place="м. Малий Ясенець", vpd="4390/26"),

    # --- Виняток: особи немає в штатці (черга рев'ю new_person) ---
    dict(id="DEMO-17", kind="leave", who="__outsider__", number="124",
         issue=-5, start=-2, end=4,
         leave_type="відпустка за сімейними обставинами",
         place="м. Кам’яна Слобода", vpd="4301/26",
         category="поза штаткою",
         note="Особи НЕМА в db/seeds/unit_roster.csv -- навмисний виняток "
              "на демонстрацію черги рев'ю (нова особа, не з реєстру)."),
]

# Які документи історії дублюються в інших форматах.
PDF_IDS = ["DEMO-03", "DEMO-04", "DEMO-07", "DEMO-09", "DEMO-13"]
PHOTO_PLAN = [("DEMO-01", "normal"), ("DEMO-02", "normal2"),
              ("DEMO-06", "png"), ("DEMO-10", "worst")]

# Резерв «не заливати в базу»: провести живцем на демо.
LIVE = [
    dict(id="LIVE-01", kind="leave", who="UNIT-0012", number="151",
         issue=-3, start=0, end=6,
         leave_type="щорічна основна відпустка за 2026 рік",
         place="м. Тихолісся", vpd="4470/26", fmt="docx"),
    dict(id="LIVE-02", kind="deployment", who="UNIT-0026", number="241",
         issue=-2, order=-3, order_number="431", start=1, end=4,
         dest="м. Кривоярськ", dest_org="військова частина Т5140",
         purpose="отримання матеріально-технічних засобів", fmt="pdf"),
    dict(id="LIVE-03", kind="leave", who="UNIT-0037", number="153",
         issue=-1, start=2, end=8,
         leave_type="відпустка у зв’язку з навчанням",
         place="м. Сухобрід", vpd="4472/26", fmt="photo"),
]

# --- масовка ---
BULK_LEAVE_TYPES = [
    "щорічна основна відпустка за 2026 рік",
    "щорічна додаткова відпустка за 2026 рік",
    "відпустка за сімейними обставинами",
    "відпустка у зв’язку з навчанням",
]
BULK_PLACES = ["м. Тихолісся", "м. Кривоярськ", "с. Соснова Гряда",
               "м. Заріччя-Долинське", "м. Сухобрід", "м. Малий Ясенець",
               "м. Кам’яна Слобода", "с. Верхня Тернівка"]
BULK_ORGS = ["військова частина Т5140", "військова частина К7719",
             "військова частина Р2268",
             "Центральна база зберігання майна (в/ч Р2268)",
             "навчальний центр (в/ч Т5140)"]
BULK_PURPOSES = ["отримання матеріально-технічних засобів",
                 "проходження курсу підвищення кваліфікації",
                 "участь у зборах командирів взводів",
                 "супроводження вантажу",
                 "передача обладнання за актом",
                 "участь у польових заняттях"]


def _mkdirs(*dirs):
    for path in dirs:
        os.makedirs(path, exist_ok=True)


def _who_of(spec, roster):
    key = spec["who"]
    if key == "__outsider__":
        return person_slots(OUTSIDER, in_roster=False)
    service_id = STORY_PEOPLE.get(key, key)
    return person_slots(roster[service_id])


def _build_one(spec, who, today):
    if spec["kind"] == "leave":
        return build_leave_values(spec, who, today)
    return build_deployment_values(spec, who, today)


def write_expected(doc_id, spec, who, v, truth):
    os.makedirs(EXPECTED_DIR, exist_ok=True)
    path = os.path.join(EXPECTED_DIR, f"{doc_id}.json")
    with open(path, "w", encoding="utf-8") as fh:
        json.dump(expected_json(doc_id, spec, who, v, truth), fh,
                  ensure_ascii=False, indent=2)
    return path


def gen_story(roster, today, *, want_pdf, want_photo, rng):
    _mkdirs(DIR_STORY, DIR_PDF, DIR_PHOTO)
    built = {}
    for spec in STORY:
        who = _who_of(spec, roster)
        v, truth = _build_one(spec, who, today)
        docx_path = os.path.join(DIR_STORY, f"{spec['id']}.docx")
        if spec["kind"] == "leave":
            fill_leave_docx(docx_path, v)
        else:
            fill_deployment_docx(docx_path, v)
        write_expected(spec["id"], spec, who, v, truth)
        built[spec["id"]] = (spec, who, v)
        print("docx", docx_path)

    if want_pdf:
        for doc_id in PDF_IDS:
            spec, who, v = built[doc_id]
            pages = (leave_pdf_pages(v) if spec["kind"] == "leave"
                     else deployment_pdf_pages(v))
            pdf_path = os.path.join(DIR_PDF, f"{doc_id}.pdf")
            render_pdf(pdf_path, pages)
            print("pdf ", pdf_path)

    if want_photo:
        for doc_id, profile in PHOTO_PLAN:
            spec, who, v = built[doc_id]
            pages = (leave_pdf_pages(v) if spec["kind"] == "leave"
                     else deployment_pdf_pages(v))
            tmp_pdf = os.path.join(DIR_PHOTO, f"{doc_id}.tmp.pdf")
            render_pdf(tmp_pdf, pages)
            out = make_photo(tmp_pdf, DIR_PHOTO, doc_id, profile, rng)
            os.remove(tmp_pdf)
            print("photo", out, f"({profile})")


def gen_live(roster, today, rng):
    _mkdirs(DIR_LIVE)
    for spec in LIVE:
        who = _who_of(spec, roster)
        v, truth = _build_one(spec, who, today)
        write_expected(spec["id"], spec, who, v, truth)
        if spec["fmt"] == "docx":
            path = os.path.join(DIR_LIVE, f"{spec['id']}.docx")
            (fill_leave_docx if spec["kind"] == "leave"
             else fill_deployment_docx)(path, v)
        else:
            pages = (leave_pdf_pages(v) if spec["kind"] == "leave"
                     else deployment_pdf_pages(v))
            pdf_path = os.path.join(DIR_LIVE, f"{spec['id']}.pdf")
            render_pdf(pdf_path, pages)
            path = pdf_path
            if spec["fmt"] == "photo":
                path = make_photo(pdf_path, DIR_LIVE, spec["id"], "normal", rng)
                os.remove(pdf_path)
        print("live", path)


def gen_bulk(roster, today, rng, count):
    """Масовка для обсягу: ті самі дві форми, інші особи штатки, дати в межах
    [сьогодні-90; сьогодні+35] (для 28.08.2026 це червень-вересень).
    ~12% документів -- із порожнім критичним полем, щоб «непідтверджених» у
    чаті не було нулем."""
    _mkdirs(DIR_BULK)
    story_ids = {STORY_PEOPLE[k] for k in STORY_PEOPLE}
    story_ids |= {s["who"] for s in LIVE}
    pool = [sid for sid in sorted(roster) if sid not in story_ids]
    rng.shuffle(pool)

    gap_every = 8  # 1 з 8 -> 12.5%
    made = 0
    for index in range(count):
        service_id = pool[index % len(pool)]
        who = person_slots(roster[service_id])
        kind = "leave" if index % 5 < 3 else "deployment"
        start_off = rng.randint(-88, 30)
        length = rng.choice([3, 4, 5, 7, 10, 12, 14, 14, 20])
        gaps = ()
        if index % gap_every == gap_every - 1:
            gaps = ((("place",), ("dates",), ("person",))[index % 3]
                    if kind == "leave" else
                    (("dest",), ("person",), ("dates",))[index % 3])
        number = str(1000 + index)
        doc_id = f"BULK-{index + 1:03d}"
        if kind == "leave":
            spec = dict(id=doc_id, kind=kind, who=service_id, number=number,
                        issue=start_off - rng.randint(2, 9),
                        start=start_off, end=start_off + length - 1,
                        ret=(start_off + length
                             if start_off + length < 0 else None),
                        leave_type=rng.choice(BULK_LEAVE_TYPES),
                        place=rng.choice(BULK_PLACES),
                        vpd=rng.choice(["не видавались", f"{rng.randint(4000, 4999)}/26"]),
                        gaps=gaps)
            v, _ = build_leave_values(spec, who, today)
            fill_leave_docx(os.path.join(DIR_BULK, f"{doc_id}.docx"), v)
        else:
            spec = dict(id=doc_id, kind=kind, who=service_id, number=number,
                        issue=start_off - rng.randint(2, 9),
                        order=start_off - rng.randint(3, 10),
                        order_number=str(rng.randint(300, 499)),
                        start=start_off, end=start_off + min(length, 7) - 1,
                        dest=rng.choice(BULK_PLACES),
                        dest_org=rng.choice(BULK_ORGS),
                        purpose=rng.choice(BULK_PURPOSES), gaps=gaps)
            v, _ = build_deployment_values(spec, who, today)
            fill_deployment_docx(os.path.join(DIR_BULK, f"{doc_id}.docx"), v)
        made += 1
    print(f"масовка: {made} docx у {DIR_BULK}")


def main(argv=None):
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--today", default=None,
                        help="дата демо у форматі YYYY-MM-DD (дефолт -- сьогодні)")
    parser.add_argument("--only", default="all",
                        help="через кому: story,pdf,photo,bulk,live,all")
    parser.add_argument("--bulk-count", type=int, default=130)
    parser.add_argument("--seed", type=int, default=20260822)
    args = parser.parse_args(argv)

    today = (dt.date.fromisoformat(args.today) if args.today
             else dt.date.today())
    parts = {p.strip() for p in args.only.split(",")}
    if "all" in parts:
        parts = {"story", "pdf", "photo", "bulk", "live"}

    roster = load_roster()
    rng = random.Random(args.seed)
    print(f"«сьогодні» = {today.isoformat()}; штатка: {len(roster)} осіб")

    if parts & {"story", "pdf", "photo"}:
        gen_story(roster, today, want_pdf="pdf" in parts,
                  want_photo="photo" in parts, rng=rng)
    if "live" in parts:
        gen_live(roster, today, rng)
    if "bulk" in parts:
        gen_bulk(roster, today, rng, args.bulk_count)
    print("еталони:", EXPECTED_DIR)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
