# -*- coding: utf-8 -*-
"""
Валідація демо-набору проти двох джерел правди — ДО прогону на сервері.

Навіщо: набір згенерований скриптом, а скрипт може згенерувати внутрішньо
суперечливий документ (днів не збігається з періодом, скасування «видане»
раніше за скасоване, особа поза штаткою там, де її бути не мусить). Такий
документ дасть у прогоні або хибний факт, або хибну статистику — і ми
припишемо це пайплайну.

Дві звірки:

  1. проти штатки (db/seeds/unit_roster.csv):
     - ПІБ, надрукований у документі, існує в реєстрі дослівно
       (прізвище на бланку друкується великими — порівнюємо без регістру);
     - звання в документі = звання цієї особи в реєстрі;
     - рівно ОДИН документ набору на особу поза штаткою (DEMO-17 за сценарієм).

  2. проти логіки (з еталонів data/eval/demo-story/per-document/*.json —
     не з самих docx: еталон і є те, що ми вважаємо правдою, а читати docx
     означало б перевіряти витяг, а не набір):
     - початок <= кінець; днів = різниця + 1; повернення >= кінця;
       дата видачі <= початку;
     - пари скасування: документ-скасувач виданий ПІЗНІШЕ за скасований і
       на ту саму особу; посилання взаємні;
     - одна особа не буває у двох ЧИННИХ періодах, що перетинаються
       (скасовані з перевірки виключені — саме для цього вони й скасовані);
     - номери документів унікальні в межах набору;
     - періоди лежать у серпні–вересні 2026 («зараз» = кінець серпня).

Вихід: перелік «підходить / не підходить і чому» + підсумок. Що не пройшло —
правиться ГЕНЕРАТОР (generate_demo_story.py) і набір перегенеровується;
еталон під результат не правиться.

Запуск (з кореня репозиторію):
    python data/eval/samples/demo-story/validate_demo_set.py
    python data/eval/samples/demo-story/validate_demo_set.py --quiet   # лише збої
Код виходу: 0 — усе підходить, 1 — є збої.
"""
import argparse
import csv
import datetime as dt
import glob
import json
import os
import re
import sys

from docx import Document

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.abspath(os.path.join(HERE, "..", "..", "..", ".."))
ROSTER_CSV = os.path.join(ROOT, "db", "seeds", "unit_roster.csv")
EXPECTED_DIR = os.path.join(ROOT, "data", "eval", "demo-story", "per-document")

DIR_STORY = os.path.join(HERE, "story")
DIR_PDF = os.path.join(HERE, "story-pdf")
DIR_PHOTO = os.path.join(HERE, "story-photo")
DIR_BULK = os.path.join(HERE, "bulk")
DIR_LIVE = os.path.join(HERE, "live")

# Вікно, у якому мусять лежати періоди набору. «Сьогодні» сценарію —
# 28.08.2026, історія дивиться на місяць уперед і на три тижні назад.
PERIOD_FROM = dt.date(2026, 8, 1)
PERIOD_TO = dt.date(2026, 9, 30)
# Дата видачі за визначенням раніша за початок, тому їй вікно ширше.
ISSUE_FROM = dt.date(2026, 7, 1)

# Рядок особи на бланку: «<звання> <ПРІЗВИЩЕ> <Ім’я> <По батькові>».
# Звання буває складене («ст. сержант», «мол. лейтенант»), тому воно тут
# нежадібним хвостом, а якорем виступає прізвище великими літерами.
UPPER = "А-ЯІЇЄҐ"
LOWER = "а-яіїєґ"
PERSON_LINE = re.compile(
    rf"^(?P<rank>[{LOWER}. ]{{3,20}}?)\s+"
    rf"(?P<surname>[{UPPER}][{UPPER}’'\-]{{2,}})\s+"
    rf"(?P<first>[{UPPER}][{LOWER}’'\-]+)\s+"
    rf"(?P<patronymic>[{UPPER}][{LOWER}’'\-]+)$"
)


class Report:
    """Збирач висновків. Кожен рядок — про один документ або одну звірку."""

    def __init__(self):
        self.rows = []          # (scope, ok, message)
        self.checks = {}        # назва звірки -> [ok, fail]

    def add(self, check, scope, ok, message=""):
        self.rows.append((scope, ok, f"{check}: {message}" if message else check))
        box = self.checks.setdefault(check, [0, 0])
        box[0 if ok else 1] += 1

    @property
    def failures(self):
        return [r for r in self.rows if not r[1]]


# ---------------------------------------------------------------------------
# Читання джерел
# ---------------------------------------------------------------------------
def load_roster():
    with open(ROSTER_CSV, encoding="utf-8") as fh:
        rows = list(csv.DictReader(fh))
    by_name = {}
    for row in rows:
        by_name[row["full_name"].casefold()] = row
    return rows, by_name


def load_expected():
    docs = {}
    for path in sorted(glob.glob(os.path.join(EXPECTED_DIR, "*.json"))):
        with open(path, encoding="utf-8") as fh:
            doc = json.load(fh)
        docs[doc["id"]] = doc
    return docs


def docx_lines(path):
    """Усі непорожні рядки документа — і з абзаців, і з комірок таблиць.
    Бланк дублює значення по злитих комірках, тому рядки унікалізуємо."""
    document = Document(path)
    seen, out = set(), []
    chunks = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    for chunk in chunks:
        for line in chunk.splitlines():
            line = " ".join(line.split())
            if line and line not in seen:
                seen.add(line)
                out.append(line)
    return out


def person_lines(lines):
    """Рядки, які виглядають як «звання ПРІЗВИЩЕ Ім’я По батькові»."""
    found = []
    for line in lines:
        match = PERSON_LINE.match(line)
        if match:
            found.append(match)
    return found


# ---------------------------------------------------------------------------
# Звірка 1: проти штатки
# ---------------------------------------------------------------------------
def check_against_roster(report, by_name, expected):
    """Для кожного docx набору: ПІБ із документа є в реєстрі дослівно,
    звання збігається. Плюс облік осіб поза штаткою."""
    outsiders = []
    empty_person = []

    files = []
    for folder, tag in ((DIR_STORY, "story"), (DIR_BULK, "bulk"),
                        (DIR_LIVE, "live")):
        for path in sorted(glob.glob(os.path.join(folder, "*.docx"))):
            files.append((tag, path))

    for tag, path in files:
        doc_id = os.path.splitext(os.path.basename(path))[0]
        scope = f"{tag}/{doc_id}"
        lines = docx_lines(path)
        matches = person_lines(lines)

        if not matches:
            # Порожнє ПІБ — навмисна прогалина масовки. Не збій набору, але
            # мусить бути видно в звіті: це майбутній needs_review.
            empty_person.append(scope)
            report.add("ПІБ у документі", scope, True,
                       "порожнє (навмисна прогалина критичного поля)")
            continue

        names = {(m.group("rank"),
                  f"{m.group('surname')} {m.group('first')} {m.group('patronymic')}")
                 for m in matches}
        if len(names) > 1:
            report.add("ПІБ у документі", scope, False,
                       f"у документі кілька різних осіб: {sorted(names)}")
            continue

        rank, printed = names.pop()
        person = by_name.get(printed.casefold())

        if person is None:
            outsiders.append((scope, f"{rank} {printed}"))
            report.add("ПІБ у штатці", scope, True,
                       f"«{printed}» у реєстрі НЕМАЄ — облік як виняток")
        else:
            report.add("ПІБ у штатці", scope, True,
                       f"«{person['full_name']}» = {person['service_id']}")
            if person["rank"] != rank:
                report.add("звання = реєстр", scope, False,
                           f"у документі «{rank}», у реєстрі «{person['rank']}»")
            else:
                report.add("звання = реєстр", scope, True, rank)

        # Еталон мусить казати про реєстр те саме, що й сам документ.
        doc = expected.get(doc_id)
        if doc is not None:
            claims_registry = doc["людина"]["є_в_реєстрі"]
            if claims_registry != (person is not None):
                report.add("еталон про реєстр", scope, False,
                           f"еталон каже є_в_реєстрі={claims_registry}, "
                           f"а в штатці {'є' if person else 'немає'}")
            else:
                report.add("еталон про реєстр", scope, True, "")

    # Виняток мусить бути рівно один — інакше це помилка генерації.
    if len(outsiders) == 1:
        report.add("особа поза штаткою — рівно одна", outsiders[0][0], True,
                   outsiders[0][1])
    else:
        report.add("особа поза штаткою — рівно одна", "набір", False,
                   f"таких документів {len(outsiders)}: "
                   f"{[o[0] for o in outsiders]} (за сценарієм мусить бути 1)")

    if empty_person:
        report.add("порожнє ПІБ — облік", "набір", True,
                   f"{len(empty_person)} документ(ів): {empty_person}")


# ---------------------------------------------------------------------------
# Звірка 2: проти логіки (еталони)
# ---------------------------------------------------------------------------
def as_date(value):
    return dt.date.fromisoformat(value) if value else None


def check_dates(report, expected):
    for doc_id, doc in sorted(expected.items()):
        truth = doc["правильні_відповіді"]
        start = as_date(truth["початок"])
        end = as_date(truth["кінець"])
        back = as_date(truth["повернення"])
        issue = as_date(truth["дата_видачі"])
        days = truth["днів"]

        if start is None or end is None:
            report.add("період заповнений", doc_id, False,
                       "у еталоні немає початку або кінця")
            continue

        report.add("початок <= кінець", doc_id, start <= end,
                   "" if start <= end else f"{start} > {end}")

        want = (end - start).days + 1
        report.add("днів = різниця + 1", doc_id, days == want,
                   "" if days == want else f"у еталоні {days}, з періоду {want}")

        if back is None:
            report.add("повернення >= кінець", doc_id, True,
                       "повернення не вписане (відкритий документ)")
        else:
            report.add("повернення >= кінець", doc_id, back >= end,
                       "" if back >= end else f"{back} < {end}")

        if issue is None:
            report.add("дата видачі <= початок", doc_id, False,
                       "дати видачі немає")
        else:
            report.add("дата видачі <= початок", doc_id, issue <= start,
                       "" if issue <= start else f"видано {issue}, початок {start}")

        in_window = PERIOD_FROM <= start <= PERIOD_TO and PERIOD_FROM <= end <= PERIOD_TO
        report.add("період у серпні–вересні 2026", doc_id, in_window,
                   "" if in_window else f"{start}..{end} поза "
                   f"{PERIOD_FROM}..{PERIOD_TO}")
        if issue is not None:
            ok = ISSUE_FROM <= issue <= PERIOD_TO
            report.add("дата видачі у вікні набору", doc_id, ok,
                       "" if ok else f"{issue} поза {ISSUE_FROM}..{PERIOD_TO}")


def check_cancellations(report, expected):
    pairs = {doc_id: doc for doc_id, doc in expected.items() if doc["пара"]}
    if not pairs:
        report.add("пари скасування знайдені", "набір", False,
                   "у еталонах немає жодної пари — сценарій вимагає дві")

    for doc_id, doc in sorted(pairs.items()):
        pair = doc["пара"]
        if pair.get("role") != "чинний":
            continue  # обробляємо пару з боку скасувача
        other_id = pair.get("replaces")
        other = expected.get(other_id)
        if other is None:
            report.add("скасований документ існує", doc_id, False,
                       f"посилання на {other_id}, якого в еталонах немає")
            continue

        scope = f"{other_id} -> {doc_id}"

        # взаємність посилань
        back_ref = (other["пара"] or {}).get("replaced_by")
        report.add("посилання пари взаємні", scope, back_ref == doc_id,
                   "" if back_ref == doc_id
                   else f"{other_id}.replaced_by = {back_ref}")

        # та сама особа
        same = doc["людина"]["service_id"] == other["людина"]["service_id"]
        report.add("скасування на ту саму особу", scope, same,
                   "" if same else f"{other['людина']['ПІБ']} != {doc['людина']['ПІБ']}")

        # скасувач виданий ПІЗНІШЕ
        new_issue = as_date(doc["правильні_відповіді"]["дата_видачі"])
        old_issue = as_date(other["правильні_відповіді"]["дата_видачі"])
        later = new_issue is not None and old_issue is not None and new_issue > old_issue
        report.add("скасувач виданий пізніше", scope, later,
                   "" if later else f"{doc_id} видано {new_issue}, "
                   f"{other_id} — {old_issue}")

        # скасований мусить бути позначений нечинним
        report.add("скасований позначений нечинним", scope,
                   other["чинний"] is False,
                   "" if other["чинний"] is False
                   else f"{other_id}.чинний = {other['чинний']}")
        report.add("скасувач позначений чинним", scope, doc["чинний"] is True,
                   "" if doc["чинний"] is True else f"{doc_id}.чинний = {doc['чинний']}")


def check_overlaps(report, expected):
    """Одна особа не буває у двох ЧИННИХ періодах, що перетинаються.
    Скасовані документи виключені — вони на те й скасовані."""
    by_person = {}
    for doc_id, doc in expected.items():
        if not doc["чинний"]:
            continue
        truth = doc["правильні_відповіді"]
        start, end = as_date(truth["початок"]), as_date(truth["кінець"])
        if start and end:
            by_person.setdefault(doc["людина"]["ПІБ"], []).append((start, end, doc_id))

    clashes = 0
    for name, periods in sorted(by_person.items()):
        periods.sort()
        for i in range(len(periods) - 1):
            a_start, a_end, a_id = periods[i]
            b_start, b_end, b_id = periods[i + 1]
            if b_start <= a_end:
                clashes += 1
                report.add("одна особа — без перетинів", f"{a_id} x {b_id}", False,
                           f"{name}: {a_start}..{a_end} перетинається з "
                           f"{b_start}..{b_end}")
        if len(periods) > 1:
            report.add("одна особа — без перетинів", name, True,
                       f"{len(periods)} чинних періодів, перетинів немає"
                       if clashes == 0 else "див. збої вище")
    return clashes


def check_numbers(report, expected):
    seen = {}
    for doc_id, doc in sorted(expected.items()):
        number = doc["правильні_відповіді"]["номер_документа"]
        seen.setdefault(number, []).append(doc_id)
    for number, ids in sorted(seen.items()):
        report.add("номер документа унікальний", f"№{number}", len(ids) == 1,
                   "" if len(ids) == 1 else f"той самий номер у {ids}")


def check_composition(report, expected):
    """Склад набору: скільки чого лежить на диску проти сценарію README."""
    wanted = {"story": (DIR_STORY, "*.docx", 17), "story-pdf": (DIR_PDF, "*.pdf", 5),
              "story-photo": (DIR_PHOTO, "*", 4), "bulk": (DIR_BULK, "*.docx", 130),
              "live": (DIR_LIVE, "*", 3)}
    for name, (folder, mask, count) in wanted.items():
        found = len([p for p in glob.glob(os.path.join(folder, mask))
                     if os.path.isfile(p)])
        report.add("склад набору", name, found == count,
                   f"{found} файл(ів)" if found == count
                   else f"{found}, а сценарій обіцяє {count}")

    report.add("еталонів на історію + резерв", "per-document",
               len(expected) == 20, f"{len(expected)} (мусить бути 20)")

    # Кожен документ історії й резерву мусить мати еталон, і навпаки.
    on_disk = {os.path.splitext(os.path.basename(p))[0]
               for p in glob.glob(os.path.join(DIR_STORY, "*.docx"))}
    on_disk |= {os.path.splitext(os.path.basename(p))[0]
                for p in glob.glob(os.path.join(DIR_LIVE, "*"))}
    missing = sorted(on_disk - set(expected))
    extra = sorted(set(expected) - on_disk)
    report.add("документ <-> еталон", "історія + резерв", not missing and not extra,
               "" if not missing and not extra
               else f"без еталона: {missing}; еталон без документа: {extra}")

    # pdf і фото мусять бути тими самими ID, що й docx історії.
    for folder, tag in ((DIR_PDF, "story-pdf"), (DIR_PHOTO, "story-photo")):
        ids = {os.path.splitext(os.path.basename(p))[0]
               for p in glob.glob(os.path.join(folder, "*")) if os.path.isfile(p)}
        strays = sorted(ids - on_disk)
        report.add("той самий ID, що в історії", tag, not strays,
                   "" if not strays else f"немає docx-відповідника: {strays}")


def check_printed_literally(report, expected):
    """Те, що еталон вважає надрукованим ПІБ, мусить бути в docx дослівно.
    Це не перевірка витягу — це перевірка, що еталон і файл не розійшлися."""
    for doc_id, doc in sorted(expected.items()):
        folder = DIR_LIVE if doc_id.startswith("LIVE") else DIR_STORY
        path = os.path.join(folder, f"{doc_id}.docx")
        if not os.path.isfile(path):
            continue  # резерв у pdf/фото — тут не читаємо
        printed = doc["надруковано"].get("PERSON_FULL", "")
        lines = docx_lines(path)
        found = any(printed and printed in line for line in lines)
        report.add("надруковане ПІБ є в docx дослівно", doc_id, found,
                   "" if found else f"«{printed}» у файлі не знайдено")


# ---------------------------------------------------------------------------
def main(argv=None):
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--quiet", action="store_true",
                        help="друкувати лише те, що не підходить")
    args = parser.parse_args(argv)

    _, by_name = load_roster()
    expected = load_expected()
    report = Report()

    check_composition(report, expected)
    check_against_roster(report, by_name, expected)
    check_printed_literally(report, expected)
    check_dates(report, expected)
    check_cancellations(report, expected)
    check_overlaps(report, expected)
    check_numbers(report, expected)

    print("== ВАЛІДАЦІЯ ДЕМО-НАБОРУ ==")
    print(f"штатка: {len(by_name)} осіб; еталонів: {len(expected)}\n")

    for scope, ok, message in report.rows:
        if ok and args.quiet:
            continue
        mark = "підходить" if ok else "НЕ ПІДХОДИТЬ"
        print(f"[{mark:>12}] {scope:<22} {message}")

    print("\n== ПІДСУМОК ПО ЗВІРКАХ ==")
    for check, (ok, fail) in report.checks.items():
        status = "усе підходить" if not fail else f"{fail} НЕ ПІДХОДИТЬ"
        print(f"{check:<38} {ok + fail:>4} перевірок — {status}")

    total = len(report.rows)
    bad = len(report.failures)
    print(f"\nусього перевірок: {total}; підходить: {total - bad}; "
          f"не підходить: {bad}")
    if bad:
        print("\nЩо робити: правити generate_demo_story.py і перегенерувати "
              "набір. Еталон під результат не правимо.")
    return 1 if bad else 0


if __name__ == "__main__":
    sys.exit(main())
